#!/usr/bin/env python3
"""Build the one-page tabular CV from the SQLite database."""

from __future__ import annotations

import argparse
import html
import json
import re
import shutil
import sqlite3
from pathlib import Path

from docx import Document

from vitamine.scripts.export_utils import compile_typst_if_available
from vitamine.paths import OUTPUT, PACKAGE, ROOT, active_db_path, output_ref

DB = active_db_path()
DEFAULT_TEMPLATE = PACKAGE / "onepage_tabular" / "ultrashort_tabular_template.docx"
DEFAULT_OUTPUT = OUTPUT / "ultrashort_tabular_cv.docx"
LANG = "en"


def connect() -> sqlite3.Connection:
    con = sqlite3.connect(DB)
    con.row_factory = sqlite3.Row
    ensure_publication_columns(con)
    return con


def ensure_publication_columns(con: sqlite3.Connection) -> None:
    table = con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='publications'").fetchone()
    if not table:
        return
    existing = {row[1] for row in con.execute("PRAGMA table_info(publications)").fetchall()}
    columns = {
        "include_short": "INTEGER NOT NULL DEFAULT 0",
        "include_ultrashort": "INTEGER NOT NULL DEFAULT 0",
        "selected_order": "INTEGER",
        "short_selected_order": "INTEGER",
        "ultrashort_selected_order": "INTEGER",
        "short_citation": "TEXT",
        "impact_factor": "REAL",
        "impact_factor_year": "TEXT",
        "metric_source": "TEXT",
        "suppress_display": "INTEGER NOT NULL DEFAULT 0",
        "quality_note": "TEXT",
    }
    for column, definition in columns.items():
        if column not in existing:
            con.execute(f"ALTER TABLE publications ADD COLUMN {column} {definition}")
    con.execute(
        """
        CREATE TABLE IF NOT EXISTS export_settings (
          profile TEXT PRIMARY KEY,
          publication_limit INTEGER NOT NULL DEFAULT 10,
          authorship_filter TEXT NOT NULL DEFAULT 'first_last'
        )
        """
    )
    con.execute(
        """
        INSERT OR IGNORE INTO export_settings (profile, publication_limit, authorship_filter)
        VALUES ('ultrashort', 10, 'first_last')
        """
    )


def clean(value: str | None) -> str:
    return (value or "").strip()


def tr(en: str, de: str) -> str:
    return de if LANG == "de" else en


def row_value(row: sqlite3.Row, field: str) -> str:
    if LANG == "de":
        german = f"{field}_de"
        if german in row.keys() and clean(row[german]):
            return clean(row[german])
    return clean(row[field])


def typ(value: str | None) -> str:
    return json.dumps(clean(value), ensure_ascii=False)


def text(value: str | None, *, bold: bool = False, size: str | None = None) -> str:
    args = []
    if bold:
        args.append('weight: "bold"')
    if size:
        args.append(f"size: {size}")
    args.append(typ(value))
    return f"#text({', '.join(args)})"


def impact_factor_label(row: sqlite3.Row) -> str:
    value = row["impact_factor"] if "impact_factor" in row.keys() else None
    if value in (None, ""):
        return ""
    try:
        formatted = f"{float(value):g}"
    except (TypeError, ValueError):
        formatted = clean(str(value))
    year_value = clean(row["impact_factor_year"] if "impact_factor_year" in row.keys() else "")
    return f"IF {formatted} ({year_value})" if year_value else f"IF {formatted}"


def citation_text(value: str | None) -> str:
    value = clean(value).strip()
    value = normalize_citation_spacing(value)
    return value.rstrip(" .")


def sentence_part(value: str | None) -> str:
    value = citation_text(value)
    return value if not value or value.endswith((".", "?", "!")) else f"{value}."


def normalize_citation_spacing(value: str | None) -> str:
    value = clean(value).strip()
    value = re.sub(r"\s+([,.;:])", r"\1", value)
    value = re.sub(r",(?=\S)", ", ", value)
    value = re.sub(r"\s{2,}", " ", value)
    return value


def typst_rich_text(value: str | None, *, bold_names: bool = False, underline: bool = False, italic: bool = False, size: str = "10.15pt") -> str:
    value = normalize_citation_spacing(value)
    if not value:
        return text("", size=size)
    if bold_names:
        pieces = []
        cursor = 0
        for match in re.finditer(r"\b(?:Andreas\s+Horn|Horn\s+A\.?|Horn)\b", value):
            if match.start() > cursor:
                before = value[cursor : match.start()]
                stripped = before.rstrip()
                if stripped:
                    pieces.append(text(stripped, size=size))
                if before and before[-1].isspace():
                    pieces.append("#h(0.28em)")
            pieces.append(text(match.group(0), bold=True, size=size))
            cursor = match.end()
        if cursor < len(value):
            after = value[cursor:]
            if after and after[0].isspace():
                pieces.append("#h(0.28em)")
                after = after.lstrip()
            if after:
                pieces.append(text(after, size=size))
        return "".join(pieces)
    body = text(value, size=size)
    if italic:
        body = f"#emph[{body}]"
    if underline:
        body = f"#underline[{body}]"
    return body


def typst_word_heading(heading: str) -> str:
    return f"#block(above: 0.12in, below: 0.15in)[{text(heading, bold=True, size='10.7pt')}]"


def typst_rule() -> str:
    return '#line(length: 100%, stroke: 0.65pt)'


def typst_rule_table(columns: str, rows: list[list[str]], *, size: str = "10.2pt") -> str:
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:]
    header_cells = ", ".join(f"[{text(cell, bold=True, size=size)}]" for cell in header)
    parts = [
        "#block(width: 100%)[",
        f"#grid(columns: ({columns}), column-gutter: 0.13in, {header_cells})",
        "#v(0.015in)",
        typst_rule(),
        "#v(0.04in)",
    ]
    for row in body:
        cells = ", ".join(f"[{text(cell, size=size)}]" for cell in row)
        parts.extend(
            [
                f"#grid(columns: ({columns}), column-gutter: 0.13in, {cells})",
                "#v(0.04in)",
            ]
        )
    parts.extend(["#v(0.004in)", typst_rule(), "]"])
    return "\n".join(parts)


def typst_awards(rows: list[str]) -> str:
    award_rows: list[str] = []
    for row in rows:
        left, sep, right = row.partition(" – ")
        if sep:
            content = text(left, bold=True, size="10.15pt") + '#h(0.16em)' + text(f"– {right}", size="10.15pt")
        else:
            content = text(row, size="10.15pt")
        award_rows.append(f"#grid(columns: (0.28in, 6.65in), column-gutter: 0.04in, [#text(\"•\")], [{content}])")
    return "\n#v(0.052in)\n".join(award_rows)


def typst_citation(row: sqlite3.Row) -> str:
    parts = []
    authors = citation_text(row["authors"])
    if authors:
        parts.append(typst_rich_text(sentence_part(authors), bold_names=True))
    title = citation_text(row["title"])
    if title:
        parts.append(text(sentence_part(title), size="10.15pt"))
    venue = citation_text(row["venue"])
    if venue:
        venue = venue.title() if venue.isupper() else venue
        parts.append(typst_rich_text(sentence_part(venue), italic=True, underline=True))
    year_value = citation_text(row["year"])
    if year_value:
        parts.append(text(sentence_part(year_value), size="10.15pt"))
    impact = impact_factor_label(row)
    if impact:
        parts.append(text(sentence_part(impact), size="10.15pt"))
    doi = citation_text(row["doi"])
    if doi:
        parts.append(text(sentence_part(f"doi:{doi}"), size="10.15pt"))
    return "#h(0.28em)".join(parts)


def typst_publications(rows: list[sqlite3.Row]) -> str:
    publication_rows: list[str] = []
    for index, row in enumerate(rows, 1):
        publication_rows.append(
            "#block(below: 0.075in)["
            f"#grid(columns: (0.38in, 6.55in), column-gutter: 0.08in, "
            f"[{text(str(index) + '.', size='10.15pt')}], [{typst_citation(row)}])"
            "]"
        )
    return "\n".join(publication_rows)


def year(value: str | None) -> str:
    text = clean(value)
    if not text:
        return ""
    slash_match = re.match(r"^\d{1,2}/\d{1,2}/(\d{2}|\d{4})$", text)
    if slash_match:
        value_int = int(slash_match.group(1))
        if value_int < 100:
            value_int += 2000 if value_int < 40 else 1900
        return str(value_int)
    match = re.search(r"\d{4}", text)
    return match.group(0) if match else text


def range_years(start: str | None, end: str | None) -> str:
    start_year = year(start)
    end_year = year(end)
    if start_year and end_year:
        return f"{start_year}–{end_year}"
    if start_year:
        return f"{start_year}–present"
    return ""


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_text(paragraph, parts: list[tuple[str, bool]]) -> None:
    clear_paragraph(paragraph)
    for text, bold in parts:
        run = paragraph.add_run(text)
        run.bold = bold


def set_simple_paragraph(paragraph, text: str, *, bold: bool = False) -> None:
    add_text(paragraph, [(text, bold)])


def split_year_prefix(text: str) -> list[tuple[str, bool]]:
    match = re.match(r"^([0-9]{4}(?:[–-][0-9]{4}|[–-]present)?)(.*)$", text)
    if not match:
        return [(text, False)]
    return [(match.group(1), True), (match.group(2), False)]


def set_cell(cell, text: str, *, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    set_simple_paragraph(paragraph, text, bold=bold)
    for extra in cell.paragraphs[1:]:
        clear_paragraph(extra)


def add_docx_piece(paragraph, value: str, *, bold_names: bool = False, italic: bool = False, underline: bool = False) -> None:
    if not value:
        return
    if not bold_names:
        run = paragraph.add_run(value)
        run.italic = italic
        run.underline = underline
        return
    cursor = 0
    for match in re.finditer(r"\b(?:Andreas\s+Horn|Horn\s+A\.?|Horn)\b", value):
        if match.start() > cursor:
            run = paragraph.add_run(value[cursor : match.start()])
            run.italic = italic
            run.underline = underline
        run = paragraph.add_run(match.group(0))
        run.bold = True
        run.italic = italic
        run.underline = underline
        cursor = match.end()
    if cursor < len(value):
        run = paragraph.add_run(value[cursor:])
        run.italic = italic
        run.underline = underline


def add_publication_docx_text(paragraph, row: sqlite3.Row) -> None:
    clear_paragraph(paragraph)
    parts: list[tuple[str, bool, bool, bool]] = []
    authors = citation_text(row["authors"])
    if authors:
        parts.append((sentence_part(authors), True, False, False))
    title = citation_text(row["title"])
    if title:
        parts.append((sentence_part(title), False, False, False))
    venue = citation_text(row["venue"])
    if venue:
        venue = venue.title() if venue.isupper() else venue
        parts.append((sentence_part(venue), False, True, True))
    year_value = citation_text(row["year"])
    if year_value:
        parts.append((sentence_part(year_value), False, False, False))
    impact = impact_factor_label(row)
    if impact:
        parts.append((sentence_part(impact), False, False, False))
    doi = citation_text(row["doi"])
    if doi:
        parts.append((sentence_part(f"doi:{doi}"), False, False, False))
    for index, (value, bold_names, italic, underline) in enumerate(parts):
        if index:
            paragraph.add_run(" ")
        add_docx_piece(paragraph, value, bold_names=bold_names, italic=italic, underline=underline)


def citation_parts(citation: str, bold_terms: tuple[str, ...]) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    cursor = 0
    pattern = re.compile("|".join(re.escape(term) for term in bold_terms if term))
    if not pattern.pattern:
        return [(citation, False)]
    for match in pattern.finditer(citation):
        if match.start() > cursor:
            parts.append((citation[cursor : match.start()], False))
        parts.append((match.group(0), True))
        cursor = match.end()
    if cursor < len(citation):
        parts.append((citation[cursor:], False))
    return parts


def parse_description_parts(row: sqlite3.Row) -> list[str]:
    return [part.strip() for part in row_value(row, "description").split("|") if part.strip()]


def row_by_title(con: sqlite3.Connection, section_key: str, title: str) -> sqlite3.Row | None:
    return con.execute(
        """
        SELECT *
        FROM cv_entries
        WHERE section_key = ? AND title = ?
        ORDER BY id DESC
        LIMIT 1
        """,
        (section_key, title),
    ).fetchone()


def education_rows(con: sqlite3.Connection) -> list[tuple[str, str, str, str]]:
    current_prof = row_by_title(con, "academic_appointments", "Professor")
    postdoc = row_by_title(con, "postdoctoral_training", "Postdoctoral Fellow")
    junior = row_by_title(con, "postdoctoral_training", "Junior Group Leader")
    phd = row_by_title(con, "education", "PhD")
    md = row_by_title(con, "education", "MD")

    rows = [(tr("Years", "Jahre"), tr("Qualification", "Qualifikation"), tr("Institution", "Institution"), tr("Field", "Fach"))]
    if current_prof:
        rows.append(
            (
                range_years(current_prof["start_date"], current_prof["end_date"]),
                row_value(current_prof, "title") or "Schilling Professor / Director",
                row_value(current_prof, "organization") or "Institute for Network Stimulation, University Hospital Cologne",
                "Computational Neurology, DBS, Connectomics",
            )
        )
    if postdoc or junior:
        rows.append(
            (
                "2016–2018",
                tr("Postdoctoral Fellow / Junior Group Leader", "Postdoctoral Fellow / Junior Group Leader"),
                "Harvard Medical School & Charité Berlin",
                "Medical Neurosciences",
            )
        )
    if phd:
        parts = parse_description_parts(phd)
        rows.append(("PhD", "PhD", row_value(phd, "organization") or (parts[2] if len(parts) > 2 else ""), parts[1] if len(parts) > 1 else "Medical Neurosciences"))
    if md:
        parts = parse_description_parts(md)
        rows.append(("MD", "MD", row_value(md, "organization") or (parts[2] if len(parts) > 2 else ""), parts[1] if len(parts) > 1 else "Medicine"))
    return rows[:5]


def position_rows(con: sqlite3.Connection) -> list[tuple[str, str]]:
    desired = [
        ("academic_appointments", "Professor", "Professor of Computational Neurology and Director, Network Stimulation Institute, University of Cologne"),
        ("hospital_appointments", "Investigator", "Investigator, Department of Stereotaxy and Functional Neurosurgery, University Hospital Cologne"),
        ("academic_appointments", "Associate Professor", "Associate Professor, Harvard Medical School"),
        ("academic_appointments", "Director, Deep Brain Stimulation Research", "Director, Deep Brain Stimulation Research, Brigham & Women's Hospital"),
        ("academic_appointments", "Director, Connectomic Neuromodulation Research", "Director, Connectomic Neuromodulation Research, Massachusetts General Hospital"),
        ("academic_appointments", "Emmy Noether Group Leader", "Emmy Noether Group Leader, Neurology, Charité Berlin"),
    ]
    rows = [(tr("Years", "Jahre"), tr("Position", "Position"))]
    for section_key, title, label in desired:
        row = row_by_title(con, section_key, title)
        if row:
            rows.append((range_years(row["start_date"], row["end_date"]), row_value(row, "title") or label))
    return rows[:7]


def award_rows(con: sqlite3.Connection) -> list[str]:
    selected = [
        ("funding", "Project C05 in 3rd funding period of CRC ELAINE", "CRC ELAINE Project C05, German Research Foundation (DFG), PI"),
        ("funding", "Schilling Professorship: Institute for Network Stimulation, Cologne", "Schilling Professorship: Institute for Network Stimulation, Schilling Foundation, PI"),
        ("honors", "International Brain Stimulation Early Career Award", "International Brain Stimulation Early Career Award"),
        ("honors", "One of the World's 'Highly Cited Researchers'", "Highly Cited Researcher, Clarivate; Falling Walls Global Call Winner"),
        ("honors", "Heinz-Maier-Leibnitz Prize", "Heinz-Maier-Leibnitz Prize, German Research Foundation"),
        ("funding", "Toward Connectomic Brain Stimulation", "DFG Emmy Noether Grant: Toward Connectomic Brain Stimulation, PI"),
    ]
    rows = []
    for section_key, title, label in selected:
        row = row_by_title(con, section_key, title)
        if row:
            if section_key == "honors":
                date_label = year(row["start_date"])
            else:
                date_label = range_years(row["start_date"], row["end_date"]) or year(row["start_date"])
            rows.append(f"{date_label} – {label}")
        elif title.startswith("One of the World's"):
            rows.append("2024 – Highly Cited Researcher, Clarivate; Falling Walls Global Call Winner")
    return rows[:6]


def publication_citation(row: sqlite3.Row) -> str:
    if clean(row["short_citation"]):
        return clean(row["short_citation"])
    authors = clean(row["authors"])
    if authors:
        parts = [part.strip() for part in authors.split(",")]
        if len(parts) > 3:
            authors = ", ".join(parts[:3]) + ", et al."
    citation_parts_out = [part for part in [authors, clean(row["title"])] if part]
    citation = ". ".join(citation_parts_out)
    if clean(row["venue"]):
        citation += f". {clean(row['venue'])}"
    if clean(row["year"]):
        citation += f". {clean(row['year'])}"
    if clean(row["doi"]):
        citation += f"; doi:{clean(row['doi'])}"
    return citation.rstrip(".") + "."


def selected_publications(con: sqlite3.Connection, limit: int) -> list[sqlite3.Row]:
    return con.execute(
        """
        SELECT *
        FROM publications
        WHERE include_ultrashort = 1
          AND COALESCE(suppress_display, 0) = 0
          AND category = 'peer_reviewed'
        ORDER BY COALESCE(ultrashort_selected_order, selected_order, 999), CAST(year AS INTEGER) DESC, id DESC
        LIMIT ?
        """,
        (limit,),
    ).fetchall()


def publication_limit(con: sqlite3.Connection, fallback: int = 10) -> int:
    row = con.execute("SELECT publication_limit FROM export_settings WHERE profile='ultrashort'").fetchone()
    return int(row["publication_limit"] if row else fallback)


def load_tabular_data(publication_limit_value: int) -> tuple[sqlite3.Row | None, list[tuple[str, str, str, str]], list[tuple[str, str]], list[str], list[sqlite3.Row]]:
    with connect() as con:
        person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
        edu = education_rows(con)
        positions = position_rows(con)
        awards = award_rows(con)
        publications = selected_publications(con, publication_limit_value)
    return person, edu, positions, awards, publications


def build_html_document(publication_limit_value: int) -> str:
    person, edu, positions, awards, publications = load_tabular_data(publication_limit_value)
    name = clean(person["display_name"] if person else "") or clean(person["full_name"] if person else "") or "Curriculum Vitae"
    title = clean(person["position_title"] if person else "")
    sections = [
        (tr("Education and Training", "Ausbildung"), [" | ".join(row) for row in edu[1:]]),
        (tr("Positions and Scientific Appointments", "Positionen und wissenschaftliche Berufungen"), [f"{a} | {b}" for a, b in positions[1:]]),
        (tr("Awards, Research Funding and Presentations", "Auszeichnungen, Forschungsförderung und Vorträge"), awards),
        (tr("Selected Publications", "Ausgewählte Publikationen"), [publication_citation(row) for row in publications]),
    ]
    body = [f"<h1>{html.escape(name)}</h1>"]
    if title:
        body.append(f"<p class=\"subtitle\">{html.escape(title)}</p>")
    for heading, rows in sections:
        body.append(f"<h2>{html.escape(heading)}</h2><ul>")
        body.extend(f"<li>{html.escape(row)}</li>" for row in rows[:publication_limit_value if heading.endswith('Publications') else len(rows)])
        body.append("</ul>")
    return f"""<!doctype html>
<html lang="{LANG}">
<head>
  <meta charset="utf-8">
  <title>Tabular One Page CV</title>
  <style>
    body {{ font-family: Arial, Helvetica, sans-serif; max-width: 850px; margin: 28px auto; color: #111; font-size: 12px; line-height: 1.32; }}
    h1 {{ font-size: 22px; margin: 0 0 4px; }}
    .subtitle {{ font-weight: 700; margin: 0 0 14px; }}
    h2 {{ font-size: 14px; margin: 14px 0 5px; border-bottom: 1px solid #999; }}
    ul {{ margin: 0 0 8px 18px; padding: 0; }}
    li {{ margin: 3px 0; }}
  </style>
</head>
<body>
{''.join(body)}
</body>
</html>
"""


def build_typst_document(publication_limit_value: int) -> str:
    person, edu, positions, awards, publications = load_tabular_data(publication_limit_value)
    name = clean(person["display_name"] if person else "") or clean(person["full_name"] if person else "") or "Curriculum Vitae"
    title = clean(person["position_title"] if person else "")
    degrees = clean(person["degrees"] if person else "")
    heading_name = f"Prof. Dr. {name}"
    if degrees:
        heading_name += f", {degrees}"
    edu_three_columns = [list(row[:3]) for row in edu]
    lines = [
        '#set page(width: 8.5in, height: 11in, margin: (left: 0.72in, right: 0.72in, top: 0.58in, bottom: 0.58in))',
        f'#set text(font: "Arial", size: 10.2pt, lang: "{LANG}", fill: black)',
        "#set par(leading: 0.78em, spacing: 0pt)",
        text(heading_name, bold=True, size="11.2pt"),
        "#linebreak()",
    ]
    if title:
        lines.append(text(title, bold=True, size="11.2pt"))
        lines.append("#linebreak()")
    lines.extend(
        [
            text("Institute for Network Stimulation & Department of Stereotaxy and Functional Neurosurgery", size="11.2pt"),
            "#linebreak()",
            text("University Hospital Cologne, Germany", size="11.2pt"),
            "#v(0.07in)",
            typst_word_heading(tr("Education and Training", "Ausbildung")),
            typst_rule_table("1.18in, 1.92in, 3.82in", edu_three_columns),
            typst_word_heading(tr("Positions and Scientific Appointments", "Positionen und wissenschaftliche Berufungen")),
            typst_rule_table("1.18in, 5.75in", [list(row) for row in positions]),
            typst_word_heading(tr("Awards, Research Funding and Presentations", "Auszeichnungen, Forschungsförderung und Vorträge")),
            typst_awards(awards),
            typst_word_heading(tr("Selected Publications", "Ausgewählte Publikationen")),
            typst_publications(publications),
        ]
    )
    return "\n".join(lines) + "\n"


def output_stem() -> str:
    return "ultrashort_tabular_cv_de" if LANG == "de" else "ultrashort_tabular_cv"


def build_all(template: Path, publication_limit_value: int, lang: str = "en") -> dict[str, str]:
    global LANG
    LANG = "de" if lang == "de" else "en"
    stem = output_stem()
    docx_path = OUTPUT / f"{stem}.docx"
    html_path = OUTPUT / f"{stem}.html"
    typ_path = OUTPUT / f"{stem}.typ"
    pdf_path = OUTPUT / f"{stem}.pdf"
    build(template, docx_path, publication_limit_value)
    html_path.write_text(build_html_document(publication_limit_value), encoding="utf-8")
    typ_path.write_text(build_typst_document(publication_limit_value), encoding="utf-8")
    pdf, warning = compile_typst_if_available(typ_path, pdf_path, ROOT)
    result = {
        "docx": f"output/{output_ref(docx_path)}",
        "html": f"output/{output_ref(html_path)}",
        "typst": f"output/{output_ref(typ_path)}",
    }
    if pdf:
        result["pdf"] = f"output/{output_ref(pdf)}"
    if warning:
        result["warning"] = warning
    return result


def build(template: Path, output: Path, publication_limit: int) -> Path:
    with connect() as con:
        person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
        edu = education_rows(con)
        positions = position_rows(con)
        awards = award_rows(con)
        publications = selected_publications(con, publication_limit)

    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template, output)
    doc = Document(output)
    doc.core_properties.author = person["display_name"] or person["full_name"]
    doc.core_properties.title = "Andreas Horn - Tabular CV"
    doc.core_properties.subject = "One-page tabular CV"

    set_simple_paragraph(doc.paragraphs[0], f"Prof. Dr. {person['display_name']}, {person['degrees']}", bold=True)
    add_text(
        doc.paragraphs[1],
        [
            (person["position_title"], True),
            ("\nInstitute for Network Stimulation & Department of Stereotaxy and Functional Neurosurgery", False),
            ("\nUniversity Hospital Cologne, Germany", False),
        ],
    )
    set_simple_paragraph(doc.paragraphs[2], "Education and Training", bold=True)
    set_simple_paragraph(doc.paragraphs[3], "Positions and Scientific Appointments", bold=True)
    set_simple_paragraph(doc.paragraphs[4], "Awards, Research Funding and Presentations", bold=True)
    set_simple_paragraph(doc.paragraphs[11], "Selected Publications", bold=True)

    education_column_count = len(doc.tables[0].columns)
    for row_idx, row in enumerate(edu):
        if education_column_count == 3:
            row_values = (row[0], row[1], row[2])
        else:
            row_values = row
        for col_idx, value in enumerate(row_values):
            set_cell(doc.tables[0].cell(row_idx, col_idx), value, bold=(row_idx == 0))
    for row_idx, row in enumerate(positions):
        for col_idx, value in enumerate(row):
            set_cell(doc.tables[1].cell(row_idx, col_idx), value, bold=(row_idx == 0))
    for paragraph, line in zip(doc.paragraphs[5:11], awards):
        add_text(paragraph, split_year_prefix(line))
    publication_paragraphs = doc.paragraphs[12 : 12 + publication_limit]
    for paragraph, publication in zip(publication_paragraphs, publications):
        add_publication_docx_text(paragraph, publication)
    for paragraph in publication_paragraphs[len(publications) :]:
        clear_paragraph(paragraph)

    doc.save(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--publication-limit", type=int)
    parser.add_argument("--lang", choices=["en", "de"], default="en")
    args = parser.parse_args()
    if args.publication_limit is None:
        with connect() as con:
            args.publication_limit = publication_limit(con)
    for name, path in build_all(args.template, args.publication_limit, args.lang).items():
        print(f"{name}: {path}")


if __name__ == "__main__":
    main()

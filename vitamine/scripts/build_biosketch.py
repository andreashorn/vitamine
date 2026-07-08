#!/usr/bin/env python3
"""Build an NIH-style biosketch from the local CV database."""

from __future__ import annotations

import html
import argparse
import json
import re
import sqlite3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from vitamine.scripts.export_utils import compile_typst_if_available
from vitamine.paths import OUTPUT, ROOT, active_db_path, output_ref

DB = active_db_path()
LANG = "en"


PERSONAL_STATEMENT = (
    "I am a medical scientist with training in neuroimaging, movement disorders, software development and both invasive and "
    "noninvasive brain stimulation. The goal of my research is to analyze and modulate brain networks to improve treatment "
    "of brain disease, predominantly in movement and psychiatric basal ganglia disorders. The primary tools I have used to "
    "pursue these goals are structural imaging and noninvasive connectivity measures derived from diffusion weighted and "
    "functional MRI. I have spent the last twelve years including a PhD focused on developing and improving methods to "
    "analyze brain stimulation sites and how their whole brain effects are mediated via distributed structural and functional "
    "brain networks. My work has been recognized with the Heinz-Maier Leibnitz Prize, which is the most prestigious scientific "
    "honor awarded for early-career researchers in Germany. I am lead developer of a scientific software that facilitates "
    "these types of analyses. The software, Lead-DBS, is distributed as open-source and has empowered academic research on "
    "all continents (>65,000 downloads, >1,000 peer-reviewed studies empowered)."
)

ONGOING_PROJECTS = [
    ("2021-2024", "Patient-specific dynamical modeling and optimization of Deep Brain Stimulation", "JPND call on Novel imaging and brain stimulation methods and technologies related to Neurodegenerative Diseases", "PI & Coordinator (multi-center grant, total: $1,335,696)", "To create a dynamic model of deep brain stimulation validated by imaging and electrophysiological data"),
    ("2021-2024", "FFOR Seed Grant Toward Personalized Circuit Therapy in OCD", "FFOR - The Foundation for OCD Research", "PI ($660,000)", "To develop a machine-learning based model for network-blending and integration between connectomics and transcriptomics in OCD."),
    ("2022-2027", "1R01 13478451 Toward Connectomic Deep Brain Stimulation in Obsessive Compulsive Disorder", "NIH", "PI ($2.5 MIO)", "To determine networks associated with optimal deep brain stimulation for OCD"),
    ("2023-2028", "UM1NS132358 BRAIN CONNECTS: The center for Large-scale Imaging of Neural Circuits", "NIH", "co-I ($23.5 MIO)", ""),
]

HIGHLIGHT_CITATIONS = [
    "Hollunder B, ... Horn A. Mapping Dysfunctional Circuits in the Frontal Cortex Using Deep Brain Stimulation. Nature Neuroscience. 2023",
    "Li N, ... Horn A. A unified connectomic target for deep brain stimulation in obsessive-compulsive disorder. Nature Communications. 2020",
    "Ríos AS, ... Horn A. Optimal deep brain stimulation sites and networks for stimulation of the fornix in Alzheimer's disease. Nature Communications. 2022",
    "Ganos C, ... Horn A. A neural network for tics: insights from causal brain lesions and deep brain stimulation. Brain. 2022",
]

POSITIONS = [
    ("2021-Present", "Associate Professor of Neurology, Harvard Medical School, Boston, MA"),
    ("2021-Present", "Director of DBS research, Center for Brain Circuit Therapeutics, Brigham & Women's Hospital Boston, MA"),
    ("2021-Present", "Director, Connectomic Neuromodulation Research, Massachusetts General Hospital, Boston, MA"),
    ("2021-Present", "Associate Scientist, Department of Neurology, Brigham & Women's Hospital Boston, MA"),
    ("2017-Present", "Emmy Noether Group Leader (Assistant Prof. equivalent), Department of Neurology, Charité - University Medicine Berlin, Germany"),
    ("2016-2017", "Research Fellowship, Harvard Medical School, Boston, MA"),
    ("2013-2015", "Research Fellowship, Department of Neurology, Charité - University Medicine Berlin, Germany"),
]

HONORS = [
    ("2022", "Heinz-Maier-Leibnitz Prize, German Research Foundation"),
    ("2020", "Data Reuse Award, BIH Quest Center for Transforming Medical Research"),
    ("2019", "Peer Review Award, top 1% in field, Publons"),
    ("2019", "Best Paper Award, Organization for Human Brain Mapping"),
    ("2018", "3 x Editor's Choice Award, Brain (Oxford Journal)"),
    ("2017", "Emmy Noether Excellence Fellowship, German Research Foundation"),
    ("2017", "Robert Koch Prize, Charité - University Medicine Berlin"),
    ("2016", "Harvard Radcliffe Institute Academic Ventures Grant"),
    ("2015", "Max Rubner Prize for Innovation, Stiftung Charité"),
]


def connect() -> sqlite3.Connection:
    con = sqlite3.connect(DB)
    con.row_factory = sqlite3.Row
    return con


def clean(value: str | None) -> str:
    return (value or "").strip()


def citation_text(value: str | None) -> str:
    value = clean(value)
    value = re.sub(r"\s+([,.;:])", r"\1", value)
    value = re.sub(r",(?=\S)", ", ", value)
    value = re.sub(r"(?<!\.)\.(?=[^\s.])", ". ", value)
    value = re.sub(r"\s{2,}", " ", value)
    return value.strip()


def set_run_font(run, *, size: float = 10.0, bold: bool = False, italic: bool = False, color: str = "111111") -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 3, line: float = 1.05) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_docx_paragraph(doc: Document, value: str, *, bold: bool = False, size: float = 10.0, after: float = 3) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after)
    run = paragraph.add_run(clean(value))
    set_run_font(run, size=size, bold=bold)


def add_docx_heading(doc: Document, value: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=8, after=3)
    run = paragraph.add_run(clean(value))
    set_run_font(run, size=11.5, bold=True, color="1f4e79")


def add_key_value(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=2)
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, bold=True)
    run = paragraph.add_run(clean(value))
    set_run_font(run)


def set_cell_text(cell, value: str, *, bold: bool = False, size: float = 9.2) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0)
    paragraph.text = ""
    run = paragraph.add_run(clean(value))
    set_run_font(run, size=size, bold=bold)


def add_two_column_table(doc: Document, rows: list[tuple[str, str]], *, left_width: float = 1.05) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(left_width)
        cells[1].width = Inches(6.25)
        set_cell_text(cells[0], left, bold=True)
        set_cell_text(cells[1], right)


def typ(value: str | None) -> str:
    return json.dumps(clean(value), ensure_ascii=False)


def text(value: str | None, *, bold: bool = False, italic: bool = False, size: str | None = None) -> str:
    args = []
    if bold:
        args.append('weight: "bold"')
    if italic:
        args.append('style: "italic"')
    if size:
        args.append(f"size: {size}")
    args.append(typ(value))
    return f"#text({', '.join(args)})"


def typst_citation(value: str | None, *, size: str = "9pt") -> str:
    value = citation_text(value)
    pieces = []
    cursor = 0
    for match in re.finditer(r"\b(?:Andreas\s+Horn|Horn\s+A\.?|Horn)\b", value):
        if match.start() > cursor:
            before = value[cursor : match.start()]
            stripped = before.rstrip()
            if stripped:
                pieces.append(text(stripped, size=size))
            if before and before[-1].isspace():
                pieces.append("#h(0.28em)")
        pieces.append(text(match.group(0), bold=True, size=size))
        cursor = match.end()
    if cursor < len(value):
        after = value[cursor:]
        if after and after[0].isspace():
            pieces.append("#h(0.28em)")
            after = after.lstrip()
        if after:
            pieces.append(text(after, size=size))
    return "".join(pieces) if pieces else text(value, size=size)


def add_docx_citation(doc: Document, value: str, *, size: float = 9.2, after: float = 1) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after)
    value = citation_text(value)
    cursor = 0
    for match in re.finditer(r"\b(?:Andreas\s+Horn|Horn\s+A\.?|Horn)\b", value):
        if match.start() > cursor:
            run = paragraph.add_run(value[cursor : match.start()])
            set_run_font(run, size=size)
        run = paragraph.add_run(match.group(0))
        set_run_font(run, size=size, bold=True)
        cursor = match.end()
    if cursor < len(value):
        run = paragraph.add_run(value[cursor:])
        set_run_font(run, size=size)


def paragraph(value: str, *, size: str | None = None) -> str:
    return f"#block(below: 0.08in)[#set par(justify: true)\n{text(value, size=size)}]\n"


def html_page(body: str) -> str:
    return f"""<!doctype html>
<html lang="{LANG}">
<head>
  <meta charset="utf-8">
  <title>NIH Biosketch</title>
  <style>
    body {{ font-family: Arial, Helvetica, sans-serif; max-width: 900px; margin: 28px auto; font-size: 12px; line-height: 1.28; color: #111; }}
    h1 {{ text-align: center; font-size: 18px; border-top: 1px solid #111; border-bottom: 1px solid #111; padding: 6px 0; }}
    h2 {{ font-size: 14px; margin-top: 18px; }}
    h3 {{ font-size: 13px; margin: 10px 0 4px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 6px 0 12px; }}
    td, th {{ border: 1px solid #777; padding: 4px 6px; vertical-align: top; }}
    .grid {{ display: grid; grid-template-columns: 110px 1fr; gap: 3px 18px; }}
    .citation {{ margin-left: 24px; }}
  </style>
</head>
<body>
{body}
</body>
</html>
"""


def contribution_rows(con: sqlite3.Connection) -> list[sqlite3.Row]:
    return con.execute(
        """
        SELECT *
        FROM biosketch_contributions
        ORDER BY ordinal, id
        """
    ).fetchall()


def build_html() -> str:
    with connect() as con:
        person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
        contributions = contribution_rows(con)
    name = "Horn, Andreas Georg, MD, PhD"
    position = "Associate Professor of Neurology"
    era = "ANHORN"
    if person:
        name = f"Horn, Andreas Georg, {clean(person['degrees']) or 'MD, PhD'}"
        position = clean(person["position_title"]) or position
        era = clean(person["era_commons"]) or era
    lines = [
        "<h1>BIOGRAPHICAL SKETCH</h1>",
        f"<p><strong>NAME:</strong> {html.escape(name)}</p>",
        f"<p><strong>eRA COMMONS USER NAME:</strong> {html.escape(era)}</p>",
        f"<p><strong>POSITION TITLE:</strong> {html.escape(position)}</p>",
        "<h2>A. Personal Statement</h2>",
        f"<p>{html.escape(PERSONAL_STATEMENT)}</p>",
        "<p>Ongoing projects that I would like to highlight include:</p>",
        '<div class="grid">',
    ]
    for dates, title, sponsor, role, purpose in ONGOING_PROJECTS:
        details = "<br>".join(html.escape(part) for part in [title, sponsor, role, purpose] if part)
        lines.extend([f"<div>{html.escape(dates)}</div>", f"<div>{details}</div>"])
    lines.extend(["</div>", "<p>Four citations that highlight my experience and qualifications for this project:</p>", "<ol>"])
    lines.extend(f"<li>{html.escape(citation_text(citation))}</li>" for citation in HIGHLIGHT_CITATIONS)
    lines.extend(["</ol>", "<h2>B. Positions, Scientific Appointments, and Honors</h2>", "<h3>Positions and Scientific Appointments</h3>", '<div class="grid">'])
    for dates, title in POSITIONS:
        lines.extend([f"<div>{html.escape(dates)}</div>", f"<div>{html.escape(title)}</div>"])
    lines.extend(["</div>", "<h3>Honors</h3>", '<div class="grid">'])
    for year, title in HONORS:
        lines.extend([f"<div>{html.escape(year)}</div>", f"<div>{html.escape(title)}</div>"])
    lines.extend(["</div>", "<h2>C. Contributions to Science</h2>"])
    for contribution in contributions:
        citations = json.loads(contribution["citations_json"] or "[]")
        lines.append(f"<p><strong>{contribution['ordinal']}. {html.escape(contribution['title'])}.</strong> {html.escape(contribution['narrative'])}</p>")
        for citation in citations:
            lines.append(f'<p class="citation">{html.escape(citation_text(citation))}</p>')
    lines.append("<p><strong>Complete List of Published Work in MyBibliography:</strong><br>https://www.ncbi.nlm.nih.gov/myncbi/andreas.horn.2/bibliography/public/</p>")
    return html_page("\n".join(lines))


def build_typst() -> str:
    with connect() as con:
        person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
        contributions = contribution_rows(con)
    name = "Horn, Andreas Georg, MD, PhD"
    position = "Associate Professor of Neurology"
    era = "ANHORN"
    if person:
        name = f"Horn, Andreas Georg, {clean(person['degrees']) or 'MD, PhD'}"
        position = clean(person["position_title"]) or position
        era = clean(person["era_commons"]) or era
    lines = [
        '#set page(width: 8.5in, height: 11in, margin: (left: 0.55in, right: 0.55in, top: 0.55in, bottom: 0.5in))',
        '#set text(font: "Arial", size: 10pt)',
        "#set par(leading: 0.42em)",
        "#line(length: 100%)",
        f"#align(center)[{text('BIOGRAPHICAL SKETCH', bold=True, size='13pt')}]",
        "#line(length: 100%)",
        "#grid(columns: (1fr), row-gutter: 0.04in,\n"
        f"  [{text('NAME:', bold=True)} {text(name)}],\n"
        f"  [{text('eRA COMMONS USER NAME:', bold=True)} {text(era)}],\n"
        f"  [{text('POSITION TITLE:', bold=True)} {text(position)}]\n"
        ")",
        f"\n{text('A. Personal Statement', bold=True, size='11pt')}",
        paragraph(PERSONAL_STATEMENT),
        text("Ongoing projects that I would like to highlight include:"),
        "#v(0.04in)",
    ]
    cells = []
    for dates, title, sponsor, role, purpose in ONGOING_PROJECTS:
        cells.append(f"[{text(dates)}]")
        cells.append(f"[{text(chr(10).join(part for part in [title, sponsor, role, purpose] if part))}]")
    lines.append("#grid(columns: (0.9in, 6.25in), gutter: 0.18in, row-gutter: 0.045in,\n" + ",\n".join(cells) + "\n)")
    lines.append(f"\n{text('Four citations that highlight my experience and qualifications for this project:')}")
    for index, citation in enumerate(HIGHLIGHT_CITATIONS, 1):
        lines.append(
            "#grid(columns: (0.25in, 6.8in), gutter: 0.08in,\n"
            f"  [{text(str(index)+'.')}],\n"
            f"  [{typst_citation(citation)}]\n"
            ")"
        )
    lines.append(f"\n#block(below: 0.06in)[{text('B. Positions, Scientific Appointments, and Honors', bold=True, size='11pt')}]")
    lines.append(f"#block(below: 0.04in)[{text('Positions and Scientific Appointments', bold=True)}]")
    cells = []
    for dates, title in POSITIONS:
        cells.append(f"[{text(dates)}]")
        cells.append(f"[{text(title)}]")
    lines.append("#grid(columns: (1.05in, 6.1in), gutter: 0.18in, row-gutter: 0.035in,\n" + ",\n".join(cells) + "\n)")
    lines.append(f"#block(above: 0.08in, below: 0.04in)[{text('Honors', bold=True)}]")
    cells = []
    for year, title in HONORS:
        cells.append(f"[{text(year)}]")
        cells.append(f"[{text(title)}]")
    lines.append("#grid(columns: (1.05in, 6.1in), gutter: 0.18in, row-gutter: 0.035in,\n" + ",\n".join(cells) + "\n)")
    lines.append(f"\n{text('C. Contributions to Science', bold=True, size='11pt')}")
    for contribution in contributions:
        citations = json.loads(contribution["citations_json"] or "[]")
        title_line = f"{contribution['ordinal']}. {contribution['title']}. "
        lines.append(f"#block(below: 0.115in)[{text(title_line, bold=True)}{text(contribution['narrative'])}]")
        for citation in citations:
            lines.append("#grid(columns: (0.22in, 6.92in), gutter: 0.08in, row-gutter: 0.015in,\n"
                         f"  [{text('')}],\n  [{typst_citation(citation)}]\n)")
    lines.append(paragraph("Complete List of Published Work in MyBibliography: https://www.ncbi.nlm.nih.gov/myncbi/andreas.horn.2/bibliography/public/", size="9pt"))
    return "\n".join(lines) + "\n"


def build_docx(path: Path) -> Path:
    with connect() as con:
        person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
        contributions = contribution_rows(con)
    name = "Horn, Andreas Georg, MD, PhD"
    position = "Associate Professor of Neurology"
    era = "ANHORN"
    if person:
        name = f"Horn, Andreas Georg, {clean(person['degrees']) or 'MD, PhD'}"
        position = clean(person["position_title"]) or position
        era = clean(person["era_commons"]) or era

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=8)
    run = title.add_run("BIOGRAPHICAL SKETCH")
    set_run_font(run, size=14, bold=True)

    add_key_value(doc, "NAME", name)
    add_key_value(doc, "eRA COMMONS USER NAME", era)
    add_key_value(doc, "POSITION TITLE", position)

    add_docx_heading(doc, "A. Personal Statement")
    add_docx_paragraph(doc, PERSONAL_STATEMENT, after=5)
    add_docx_paragraph(doc, "Ongoing projects that I would like to highlight include:", after=3)
    add_two_column_table(
        doc,
        [(dates, "\n".join(part for part in [title, sponsor, role, purpose] if part)) for dates, title, sponsor, role, purpose in ONGOING_PROJECTS],
        left_width=1.0,
    )
    add_docx_paragraph(doc, "Four citations that highlight my experience and qualifications for this project:", after=2)
    for index, citation in enumerate(HIGHLIGHT_CITATIONS, 1):
        add_docx_citation(doc, f"{index}. {citation}", size=9.2, after=1)

    add_docx_heading(doc, "B. Positions, Scientific Appointments, and Honors")
    add_docx_paragraph(doc, "Positions and Scientific Appointments", bold=True, after=2)
    add_two_column_table(doc, POSITIONS, left_width=1.1)
    add_docx_paragraph(doc, "Honors", bold=True, after=2)
    add_two_column_table(doc, HONORS, left_width=1.1)

    add_docx_heading(doc, "C. Contributions to Science")
    for contribution in contributions:
        citations = json.loads(contribution["citations_json"] or "[]")
        add_docx_paragraph(doc, f"{contribution['ordinal']}. {contribution['title']}. {contribution['narrative']}", bold=True, after=2)
        for citation in citations:
            add_docx_citation(doc, citation, size=9.2, after=1)
    add_docx_paragraph(
        doc,
        "Complete List of Published Work in MyBibliography: https://www.ncbi.nlm.nih.gov/myncbi/andreas.horn.2/bibliography/public/",
        size=9.2,
    )

    doc.core_properties.title = "NIH Biosketch Draft"
    doc.core_properties.author = name
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def output_stem() -> str:
    return "biosketch_de" if LANG == "de" else "biosketch"


def build(lang: str = "en") -> dict[str, str]:
    global LANG
    LANG = "de" if lang == "de" else "en"
    OUTPUT.mkdir(parents=True, exist_ok=True)
    stem = output_stem()
    html_path = OUTPUT / f"{stem}.html"
    typ_path = OUTPUT / f"{stem}.typ"
    pdf_path = OUTPUT / f"{stem}.pdf"
    docx_path = OUTPUT / f"{stem}.docx"
    html_path.write_text(build_html(), encoding="utf-8")
    typ_path.write_text(build_typst(), encoding="utf-8")
    pdf, warning = compile_typst_if_available(typ_path, pdf_path, ROOT)
    docx = build_docx(docx_path)
    result = {
        "html": f"output/{output_ref(html_path)}",
        "typst": f"output/{output_ref(typ_path)}",
    }
    if pdf:
        result["pdf"] = f"output/{output_ref(pdf)}"
    if docx:
        result["docx"] = f"output/{output_ref(docx)}"
    warnings = [item for item in (warning,) if item]
    if warnings:
        result["warning"] = " ".join(warnings)
    return result


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--lang", choices=["en", "de"], default="en")
    args = parser.parse_args()
    for name, path in build(args.lang).items():
        print(f"{name}: {path}")

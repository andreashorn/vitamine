#!/usr/bin/env python3
"""Import an uploaded CV into the active VitaMine database."""

from __future__ import annotations

import datetime as dt
import json
import os
import re
import shutil
import socket
import sqlite3
import subprocess
import time
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any

from docx import Document

from vitamine.paths import APP_SUPPORT, ROOT, bundled_model_path, tool_path
from vitamine.scripts.import_background_docs import (
    PUBLICATION_SECTIONS,
    clean_markup,
    extract_biosketch_contributions,
    extract_entries,
    normalize_citation,
    split_sections,
)


PERSON_FIELDS = [
    "full_name",
    "display_name",
    "degrees",
    "position_title",
    "office_address",
    "home_address",
    "work_phone",
    "work_email",
    "place_of_birth",
    "era_commons",
    "orcid_id",
]

ENTRY_FIELDS = [
    "section_key",
    "subcategory",
    "start_date",
    "end_date",
    "title",
    "organization",
    "location",
    "role",
    "amount",
    "description",
    "raw_text",
    "confidence",
    "include_extended",
    "include_long",
    "include_short",
    "include_biosketch",
    "language",
    "source_note",
]

PUBLICATION_FIELDS = [
    "item_type",
    "category",
    "authors",
    "title",
    "venue",
    "year",
    "doi",
    "pmid",
    "url",
    "abstract",
    "extra",
    "raw_citation",
    "confidence",
    "include_short",
    "include_ultrashort",
    "short_citation",
    "suppress_display",
    "quality_note",
]

SECTION_KEYS = {
    "education": "Education",
    "postdoctoral_training": "Postdoctoral Training",
    "academic_appointments": "Faculty Academic Appointments",
    "hospital_appointments": "Hospital / Affiliated Appointments",
    "professional_positions": "Other Professional Positions",
    "committee_service": "Committee Service",
    "professional_societies": "Professional Societies",
    "grant_review": "Grant Review Activities",
    "editorial_activities": "Editorial Activities",
    "honors": "Honors and Prizes",
    "funding": "Research Funding",
    "teaching": "Teaching",
    "mentoring": "Trainees and Their Successes",
    "invited_presentations": "Invited Teaching and Presentations",
    "clinical_activities": "Clinical Activities and Innovations",
    "education_innovations": "Teaching and Education Innovations",
    "community_service": "Community Service",
}

PERSON_NAME_STOPLIST = {
    "curriculum vitae",
    "cv",
    "biosketch",
    "biographical sketch",
    "personal statement",
    "nih biosketch",
}

PUBLICATION_HINTS = (
    "publication",
    "bibliography",
    "scholarship",
    "peer-reviewed",
    "peer reviewed",
    "journal articles",
    "abstracts",
    "preprints",
    "manuscripts in preparation",
    "manuscripts under review",
    "poster presentations",
    "posters",
    "patents",
)

HEADING_HINTS = {
    "education": ("education", "training", "medical school", "graduate"),
    "postdoctoral_training": ("postdoctoral", "post-doctoral", "residency", "fellowship"),
    "academic_appointments": ("academic appointment", "faculty appointment", "appointment"),
    "hospital_appointments": ("hospital", "clinical appointment"),
    "professional_positions": ("employment", "professional position", "experience", "positions"),
    "committee_service": ("committee", "service"),
    "professional_societies": ("societies", "memberships", "professional society"),
    "grant_review": ("grant review", "study section", "review activities"),
    "editorial_activities": ("editorial", "reviewer"),
    "honors": ("honors", "awards", "prizes", "recognition"),
    "funding": ("funding", "grants", "research support"),
    "teaching": ("teaching", "courses", "lectures"),
    "mentoring": ("mentoring", "supervision", "trainees"),
    "invited_presentations": ("presentations", "invited talks", "lectureships"),
    "clinical_activities": ("clinical activities", "clinical innovations"),
    "education_innovations": ("education innovations", "teaching innovations"),
    "community_service": ("community service", "outreach"),
}


def _text(value: Any) -> str:
    return str(value or "").strip()


def _clean_identifier(value: Any) -> str:
    return re.sub(r"\s+", " ", _text(value)).strip(" ,;:.")


def _looks_like_document_title(value: str) -> bool:
    clean = _clean_identifier(value).casefold()
    return clean in PERSON_NAME_STOPLIST or clean.startswith("curriculum vitae") or clean.endswith("biosketch")


def clean_person_name(value: Any) -> str:
    name = _clean_identifier(value)
    name = re.sub(r"^(full\s+name|display\s+name|name)\s*:\s*", "", name, flags=re.I).strip()
    if "," in name:
        parts = [part.strip() for part in name.split(",") if part.strip()]
        degree_like = {"md", "phd", "msc", "ma", "ba", "dr", "dr.", "prof", "prof."}
        if len(parts) >= 2 and all(part.casefold().replace(".", "") in degree_like for part in parts[1:]):
            name = parts[0]
        elif len(parts) >= 2 and parts[1].casefold() not in degree_like:
            name = f"{parts[1]} {parts[0]}"
    return _clean_identifier(name)


def run_pandoc(path: Path) -> str | None:
    pandoc = tool_path("pandoc") or shutil.which("pandoc")
    if not pandoc:
        return None
    return subprocess.check_output([str(pandoc), "-t", "markdown", "--wrap=none", str(path)], cwd=ROOT, text=True)


def docx_to_text(path: Path) -> str:
    markdown = run_pandoc(path)
    if markdown:
        return markdown
    doc = Document(path)
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def pdf_to_text(path: Path) -> str:
    pdftotext = tool_path("pdftotext")
    if pdftotext:
        return subprocess.check_output([pdftotext, "-layout", str(path), "-"], text=True, errors="replace")
    try:
        import fitz  # type: ignore

        with fitz.open(path) as doc:
            return "\n".join(page.get_text("text") for page in doc)
    except Exception:
        pass
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise RuntimeError("Could not extract PDF text. Install Poppler, PyMuPDF, or pypdf for PDF CV import.") from exc


def extract_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return docx_to_text(path), "docx"
    if suffix == ".pdf":
        return pdf_to_text(path), "pdf"
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace"), suffix.lstrip(".")
    raise RuntimeError("Please upload a DOCX, PDF, TXT, or Markdown CV.")


def classify_heading(line: str) -> str | None:
    clean = clean_markup(line).strip(":- ").casefold()
    if not clean or len(clean) > 90:
        return None
    if any(hint in clean for hint in PUBLICATION_HINTS):
        return "publications"
    for key, hints in HEADING_HINTS.items():
        if any(hint in clean for hint in hints):
            return key
    return None


def generic_sections(text: str) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current = {"section_key": "document_header", "title": "Document Header", "raw_markdown": "", "ordinal": 0}
    lines: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        key = classify_heading(line)
        looks_like_heading = bool(key) and not re.search(r"\b(19\d{2}|20\d{2})\b", line)
        if looks_like_heading:
            if lines:
                current["raw_markdown"] = "\n".join(lines).strip()
                sections.append(current)
            current = {
                "section_key": key,
                "title": clean_markup(line),
                "raw_markdown": "",
                "ordinal": len(sections),
            }
            lines = []
        else:
            lines.append(raw)
    if lines:
        current["raw_markdown"] = "\n".join(lines).strip()
        sections.append(current)
    return sections


def heuristic_person(text: str) -> dict[str, str]:
    lines = [clean_markup(line) for line in text.splitlines() if clean_markup(line)]
    person: dict[str, str] = {}
    for line in lines[:30]:
        if "@" in line and not person.get("work_email"):
            match = re.search(r"[\w.+-]+@[\w.-]+\.\w+", line)
            if match:
                person["work_email"] = match.group(0)
        if "orcid" in line.casefold() and not person.get("orcid_id"):
            match = re.search(r"\d{4}-\d{4}-\d{4}-[\dX]{4}", line)
            if match:
                person["orcid_id"] = match.group(0)
    for line in lines[:8]:
        if _looks_like_document_title(line):
            continue
        if len(line) < 80 and re.search(r"[A-Za-z]", line) and not any(char.isdigit() for char in line):
            if len(line.split()) in {2, 3, 4, 5}:
                person["full_name"] = line
                person["display_name"] = line
                break
    return person


def heuristic_entries(text: str) -> tuple[list[dict[str, Any]], int]:
    sections = split_sections(text)
    if len(sections) <= 1:
        sections = generic_sections(text)
    entries: list[dict[str, Any]] = []
    skipped_publication_sections = 0
    for section in sections:
        if section["section_key"] == "publications":
            skipped_publication_sections += 1
            continue
        entries.extend(extract_entries(section))
    for entry in entries:
        entry.setdefault("include_extended", 1)
        entry.setdefault("include_long", 1)
        entry.setdefault("include_short", 1 if entry.get("section_key") in {"education", "honors", "academic_appointments"} else 0)
        entry.setdefault("include_biosketch", 1 if entry.get("section_key") == "honors" else 0)
        entry.setdefault("language", "en")
        entry.setdefault("source_note", "Imported from uploaded CV by heuristic parser.")
    return entries, skipped_publication_sections


def publication_sections(text: str) -> list[dict[str, Any]]:
    sections = split_sections(text)
    if len(sections) <= 1:
        sections = generic_sections(text)
    return [section for section in sections if section.get("section_key") in PUBLICATION_SECTIONS]


def document_sections(text: str) -> list[dict[str, Any]]:
    sections = split_sections(text)
    return sections if len(sections) > 1 else generic_sections(text)


def coalesce_numbered_citations(raw_markdown: str) -> list[str]:
    citations: list[str] = []
    current: list[str] = []
    expected_ordinal: int | None = 1
    for raw_line in raw_markdown.splitlines():
        text = clean_markup(raw_line)
        if not text:
            continue
        ordinal_match = re.match(r"^(\d+)\.\s+", text)
        if ordinal_match and (expected_ordinal is None or int(ordinal_match.group(1)) == expected_ordinal):
            if current:
                citations.append(" ".join(current))
            current = [text]
            expected_ordinal = int(ordinal_match.group(1)) + 1
        elif current:
            current.append(text)
    if current:
        citations.append(" ".join(current))
    return citations


def heuristic_publications(text: str) -> list[dict[str, Any]]:
    publications: list[dict[str, Any]] = []
    for section in publication_sections(text):
        category = PUBLICATION_SECTIONS.get(str(section.get("section_key") or ""), "other")
        for citation in coalesce_numbered_citations(str(section.get("raw_markdown") or "")):
            pub = normalize_citation(citation)
            if not pub.get("raw_citation"):
                continue
            pub["category"] = category
            pub["item_type"] = "patent" if category == "patents" else "journal-article"
            pub["confidence"] = "medium"
            publications.append(repair_publication_metadata(pub))
    return publications


def publication_from_citation(raw_citation: str, category: str = "peer_reviewed") -> dict[str, Any] | None:
    citation = clean_markup(raw_citation)
    if not citation:
        return None
    pub = normalize_citation(citation)
    if not pub.get("raw_citation"):
        return None
    pub["category"] = category
    pub["item_type"] = "journal-article"
    pub["confidence"] = "medium"
    return repair_publication_metadata(pub)


def doi_from_citation(raw_citation: str) -> str:
    match = re.search(r"(?:doi:\s*|https?://(?:dx\.)?doi\.org/)(10\.[^\s,;]+)", raw_citation, flags=re.I)
    if match:
        return match.group(1).rstrip(").")
    match = re.search(r"\b(10\.\d{4,9}/[^\s,;]+)", raw_citation, flags=re.I)
    return match.group(1).rstrip(").") if match else ""


def repair_publication_metadata(publication: dict[str, Any]) -> dict[str, Any]:
    raw_citation = _text(publication.get("raw_citation"))
    if not raw_citation:
        return publication
    doi = doi_from_citation(raw_citation)
    if doi:
        publication["doi"] = doi
    year_match = re.search(r"\((19\d{2}|20\d{2})\)", raw_citation) or re.search(r"\b(19\d{2}|20\d{2})\b", raw_citation)
    if year_match:
        publication["year"] = year_match.group(1)
    apa_match = re.match(r"^(?:\d+\.\s*)?(?P<authors>.+?)\s+\((?P<year>19\d{2}|20\d{2})\)\.?\s+(?P<rest>.+)$", raw_citation)
    if apa_match:
        publication["authors"] = apa_match.group("authors").strip(" .")
        publication["year"] = apa_match.group("year")
        rest = re.sub(r"\s+(?:https?://(?:dx\.)?doi\.org/|doi:)\S+.*$", "", apa_match.group("rest"), flags=re.I).strip()
        title_match = re.match(r"(?P<title>.+?)\.\s+(?P<venue>[^.]+?)(?:,\s|\.\s|$)", rest)
        if title_match:
            publication["title"] = title_match.group("title").strip()
            publication["venue"] = title_match.group("venue").strip(" .")
    return publication


def heuristic_contributions(text: str) -> list[dict[str, Any]]:
    contributions: list[dict[str, Any]] = []
    for section in document_sections(text):
        extracted = extract_biosketch_contributions(section)
        contributions.extend(extracted)
        if extracted or section.get("section_key") != "biosketch_contributions":
            continue
        current: dict[str, Any] | None = None
        current_citation: dict[str, str] | None = None
        expected_ordinal = 1
        for raw_line in str(section.get("raw_markdown") or "").splitlines():
            line = clean_markup(raw_line)
            if not line:
                continue
            contribution_match = re.match(r"^(\d+)\.\s+(.+)$", line)
            if contribution_match and int(contribution_match.group(1)) == expected_ordinal:
                if current:
                    contributions.append(current)
                expected_ordinal += 1
                body = contribution_match.group(2).strip()
                title, _, rest = body.partition(". ")
                current = {
                    "ordinal": int(contribution_match.group(1)),
                    "title": title.strip(". "),
                    "narrative": rest.strip(),
                    "raw_text": body,
                    "citations": [],
                    "confidence": "medium",
                }
                current_citation = None
                continue
            if not current:
                continue
            citation_match = re.match(r"^([a-d])\.\s+(.+)$", line, flags=re.I)
            if citation_match:
                current_citation = {
                    "label": citation_match.group(1).lower(),
                    "raw_citation": citation_match.group(2).strip(),
                    "pmid": "",
                    "doi": doi_from_citation(citation_match.group(2)),
                }
                current["citations"].append(current_citation)
                continue
            if current_citation:
                current_citation["raw_citation"] = f"{current_citation['raw_citation']} {line}".strip()
                current_citation["doi"] = current_citation.get("doi") or doi_from_citation(current_citation["raw_citation"])
            else:
                current["narrative"] = f"{current.get('narrative', '')} {line}".strip()
                current["raw_text"] = f"{current.get('raw_text', '')}\n{line}".strip()
        if current:
            contributions.append(current)
    return contributions


def heuristic_narrative_report(text: str) -> dict[str, str] | None:
    for section in document_sections(text):
        key = str(section.get("section_key") or "")
        title = _text(section.get("title"))
        body = "\n".join(clean_markup(line) for line in str(section.get("raw_markdown") or "").splitlines() if clean_markup(line))
        if key in {"biosketch_personal_statement", "narrative_report"} and body:
            return {"title": title or "Narrative Report", "body": body, "title_de": "Narrativer Bericht", "body_de": ""}
        if "Narrative Report" in body:
            _, _, narrative = body.partition("Narrative Report")
            narrative = narrative.strip()
            if narrative:
                return {"title": "Narrative Report", "body": narrative, "title_de": "Narrativer Bericht", "body_de": ""}
    return None


def contribution_publications(contributions: list[dict[str, Any]]) -> list[dict[str, Any]]:
    publications: list[dict[str, Any]] = []
    for contribution in contributions:
        for citation in contribution.get("citations") or []:
            if not isinstance(citation, dict):
                continue
            pub = publication_from_citation(_text(citation.get("raw_citation")))
            if pub:
                publications.append(pub)
    return publications


def compact_llm_text(text: str, max_chars: int = 11000) -> str:
    sections = document_sections(text)
    chunks: list[str] = []
    header = "\n".join(line for line in text.splitlines()[:80] if clean_markup(line))
    if header:
        chunks.append("Document Header\n" + header[:2500])
    priority = {
        "biosketch_personal_statement",
        "biosketch_contributions",
        "narrative_report",
        "document_header",
        "education",
        "academic_appointments",
        "professional_positions",
    }
    for section in sections:
        key = str(section.get("section_key") or "")
        if key not in priority:
            continue
        chunk = f"{section.get('title') or key}\n{section.get('raw_markdown') or ''}".strip()
        if chunk:
            chunks.append(chunk[:5000])
    compact = "\n\n---\n\n".join(chunks)
    return compact[:max_chars] if compact else text[:max_chars]


def cv_json_schema() -> dict[str, Any]:
    entry_properties = {
        field: {"type": ["string", "null"]}
        for field in [
            "section_key",
            "subcategory",
            "start_date",
            "end_date",
            "title",
            "organization",
            "location",
            "role",
            "amount",
            "description",
            "raw_text",
            "confidence",
            "language",
        ]
    }
    publication_properties = {
        field: {"type": ["string", "null"]}
        for field in [
            "item_type",
            "category",
            "authors",
            "title",
            "venue",
            "year",
            "doi",
            "pmid",
            "url",
            "abstract",
            "extra",
            "raw_citation",
            "confidence",
            "short_citation",
            "quality_note",
        ]
    }
    citation_properties = {
        "label": {"type": ["string", "null"]},
        "raw_citation": {"type": ["string", "null"]},
        "pmid": {"type": ["string", "null"]},
        "doi": {"type": ["string", "null"]},
    }
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "person": {
                "type": "object",
                "additionalProperties": False,
                "properties": {field: {"type": ["string", "null"]} for field in PERSON_FIELDS},
                "required": PERSON_FIELDS,
            },
            "entries": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": entry_properties,
                    "required": list(entry_properties.keys()),
                },
            },
            "skipped_publication_count": {"type": "integer"},
            "contributions": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "title": {"type": ["string", "null"]},
                        "narrative": {"type": ["string", "null"]},
                        "raw_text": {"type": ["string", "null"]},
                        "citations": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": citation_properties,
                                "required": list(citation_properties.keys()),
                            },
                        },
                        "confidence": {"type": ["string", "null"]},
                    },
                    "required": ["title", "narrative", "raw_text", "citations", "confidence"],
                },
            },
            "publications": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": publication_properties,
                    "required": list(publication_properties.keys()),
                },
            },
            "narrative_report": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "title": {"type": ["string", "null"]},
                    "body": {"type": ["string", "null"]},
                    "title_de": {"type": ["string", "null"]},
                    "body_de": {"type": ["string", "null"]},
                },
                "required": ["title", "body", "title_de", "body_de"],
            },
            "warnings": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["person", "entries", "skipped_publication_count", "contributions", "publications", "narrative_report", "warnings"],
    }


def llm_prompt(text: str) -> str:
    labels = "\n".join(f"- {key}: {label}" for key, label in SECTION_KEYS.items())
    return f"""Extract this academic CV into VitaMine database fields.

Return one valid JSON object only, with these top-level keys:
person, entries, skipped_publication_count, contributions, publications, narrative_report, warnings.
Use empty arrays or null-valued fields when content is not present.

Rules:
- Return only data that is clearly present in the CV.
- The person is the researcher described by the document, never a document title such as "Curriculum Vitae" or "Biosketch".
- Put degrees such as MD, PhD, Dr. med., MSc in degrees, not in full_name unless the source prints no cleaner name.
- Use these section_key values only:
{labels}
- Do not import publications as entries. Instead, parse publication-looking citations into publications whenever enough citation text is present.
- For publications, keep the original citation in raw_citation, split authors/title/venue/year/doi/pmid when possible, set category to peer_reviewed, patents, books_chapters, preprints, manuscripts_in_preparation, poster_presentations, abstract, review, or other, and set confidence to high, medium, or low.
- If the CV contains NIH-style Contributions to Science, extract each contribution title and narrative into contributions. Put cited products in each contribution's citations array and also include them in publications if enough citation text is present.
- If the document contains a narrative report, personal statement, summary statement, or biosketch narrative paragraph, copy it into narrative_report.body. Use body_de only for German text.
- Dates may be years or date ranges as printed in the CV.
- raw_text should preserve the source line/block for each entry.
- confidence should be high, medium, or low.
- language should usually be en unless the CV content is German.

CV text:
{text[:60000]}
"""


def parse_json_response(data: bytes) -> dict[str, Any]:
    payload = json.loads(data.decode("utf-8"))
    content = ""
    if "message" in payload:
        content = payload.get("message", {}).get("content", "")
    elif "choices" in payload:
        content = payload["choices"][0].get("message", {}).get("content", "")
    else:
        return payload
    if isinstance(content, dict):
        return content
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?\s*", "", content, flags=re.I)
        content = re.sub(r"\s*```$", "", content).strip()
    return json.loads(content)


def merge_llm_objects(results: list[dict[str, Any]]) -> dict[str, Any]:
    merged: dict[str, Any] = {
        "person": {},
        "entries": [],
        "skipped_publication_count": 0,
        "contributions": [],
        "publications": [],
        "narrative_report": {"title": None, "body": None, "title_de": None, "body_de": None},
        "warnings": [],
    }
    seen_entries: set[tuple[str, str]] = set()
    seen_publications: set[tuple[str, str, str]] = set()
    seen_contributions: set[tuple[str, str]] = set()
    narrative_bodies: list[str] = []
    for result in results:
        person = result.get("person")
        if isinstance(person, dict):
            for field in PERSON_FIELDS:
                value = _text(person.get(field))
                if value and not _text(merged["person"].get(field)):
                    merged["person"][field] = value
        for entry in result.get("entries") or []:
            if not isinstance(entry, dict):
                continue
            key = (_text(entry.get("section_key")), _text(entry.get("raw_text") or entry.get("title") or entry.get("description")).casefold())
            if key[1] and key not in seen_entries:
                seen_entries.add(key)
                merged["entries"].append(entry)
        for publication in result.get("publications") or []:
            if not isinstance(publication, dict):
                continue
            normalized = normalize_publication(publication)
            if not normalized:
                continue
            key = (normalized["doi"].casefold(), normalized["pmid"].casefold(), normalized["raw_citation"].casefold())
            if key not in seen_publications:
                seen_publications.add(key)
                merged["publications"].append(publication)
        for contribution in result.get("contributions") or []:
            if not isinstance(contribution, dict):
                continue
            key = (_text(contribution.get("title")).casefold(), _text(contribution.get("narrative") or contribution.get("raw_text")).casefold())
            if key[0] and key[1] and key not in seen_contributions:
                seen_contributions.add(key)
                merged["contributions"].append(contribution)
        report = result.get("narrative_report")
        if isinstance(report, dict):
            title = _text(report.get("title"))
            body = _text(report.get("body"))
            title_de = _text(report.get("title_de"))
            body_de = _text(report.get("body_de"))
            if title and not _text(merged["narrative_report"].get("title")):
                merged["narrative_report"]["title"] = title
            if body and body not in narrative_bodies:
                narrative_bodies.append(body)
            if title_de and not _text(merged["narrative_report"].get("title_de")):
                merged["narrative_report"]["title_de"] = title_de
            if body_de and not _text(merged["narrative_report"].get("body_de")):
                merged["narrative_report"]["body_de"] = body_de
        try:
            merged["skipped_publication_count"] += int(result.get("skipped_publication_count") or 0)
        except (TypeError, ValueError):
            pass
        warnings = result.get("warnings") or []
        if isinstance(warnings, list):
            merged["warnings"].extend(str(warning) for warning in warnings if warning)
    if narrative_bodies:
        merged["narrative_report"]["body"] = "\n\n".join(narrative_bodies)
    return merged


def llm_text_chunks(text: str, max_chars: int = 8500) -> list[str]:
    chunks: list[str] = []
    sections = document_sections(text)
    if not sections:
        sections = [{"title": "Document", "section_key": "document", "raw_markdown": text}]
    current_chunk = ""
    for section in sections:
        title = _text(section.get("title") or section.get("section_key") or "Section")
        raw = str(section.get("raw_markdown") or "")
        prefix = f"Section: {title}\n"
        current = prefix
        for line in raw.splitlines():
            addition = f"{line}\n"
            if len(current) + len(addition) > max_chars and current.strip() != prefix.strip():
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = ""
                chunks.append(current.strip())
                current = prefix + addition
            else:
                current += addition
        section_text = current.strip()
        if section_text and section_text != prefix.strip():
            if current_chunk and len(current_chunk) + len(section_text) + 2 > max_chars:
                chunks.append(current_chunk.strip())
                current_chunk = section_text
            else:
                current_chunk = f"{current_chunk}\n\n{section_text}".strip()
    if current_chunk:
        chunks.append(current_chunk.strip())
    if not chunks:
        chunks = [text[:max_chars]]
    return chunks


def call_ollama(text: str, settings: dict[str, str]) -> dict[str, Any]:
    base = (settings.get("ollama_url") or "http://127.0.0.1:11434").rstrip("/")
    model = settings.get("ollama_model") or "llama3.1:8b"
    body = {
        "model": model,
        "messages": [{"role": "user", "content": llm_prompt(text)}],
        "stream": False,
        "format": cv_json_schema(),
        "options": {"temperature": 0},
    }
    request = urllib.request.Request(
        f"{base}/api/chat",
        data=json.dumps(body).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=180) as response:
        return parse_json_response(response.read())


def call_openai_compatible(text: str, settings: dict[str, str], strict_schema: bool = True) -> dict[str, Any]:
    base = (settings.get("api_base_url") or "https://api.openai.com/v1").rstrip("/")
    model = settings.get("api_model") or "gpt-4.1-mini"
    api_key = settings.get("api_key") or ""
    response_format: dict[str, Any] = {"type": "json_object"}
    if strict_schema:
        response_format = {
            "type": "json_schema",
            "json_schema": {"name": "vitamine_cv_import", "strict": True, "schema": cv_json_schema()},
        }
    body = {
        "model": model,
        "messages": [{"role": "user", "content": llm_prompt(text)}],
        "temperature": 0,
        "max_tokens": int(settings.get("api_max_tokens") or 4096),
        "response_format": response_format,
    }
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    request = urllib.request.Request(
        f"{base}/chat/completions",
        data=json.dumps(body).encode("utf-8"),
        headers=headers,
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=180) as response:
        return parse_json_response(response.read())


def free_local_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def wait_for_llama_server(base_url: str, process: subprocess.Popen, timeout: float = 240.0) -> None:
    deadline = time.time() + timeout
    last_error = ""
    while time.time() < deadline:
        if process.poll() is not None:
            stderr = ""
            if process.stderr:
                try:
                    stderr = process.stderr.read()[-1200:]
                except OSError:
                    stderr = ""
            detail = f": {stderr.strip()}" if stderr.strip() else ""
            raise RuntimeError(f"Bundled llama-server exited before it was ready (code {process.returncode}){detail}.")
        try:
            with urllib.request.urlopen(f"{base_url}/health", timeout=2) as response:
                if response.status < 500:
                    return
        except Exception as exc:
            last_error = str(exc)
        time.sleep(0.5)
    raise RuntimeError(f"Bundled llama-server did not become ready: {last_error}")


def call_bundled_llama(text: str, settings: dict[str, str]) -> dict[str, Any]:
    server = tool_path("llama-server")
    if not server:
        raise RuntimeError("Bundled llama-server was not found. Run scripts/install_export_tools.py --tool llama-server before packaging.")
    server = str(cached_runtime_tool(Path(server)))
    configured_model = _text(settings.get("bundled_llama_model_path"))
    model = Path(configured_model).expanduser() if configured_model else bundled_model_path()
    if not model or not model.exists():
        raise RuntimeError("Bundled local LLM model was not found. Add vendor/models/vitamine-import.gguf before packaging.")
    model = cached_runtime_model(model)
    port = free_local_port()
    base_url = f"http://127.0.0.1:{port}"
    command = [
        server,
        "--host",
        "127.0.0.1",
        "--port",
        str(port),
        "--model",
        str(model),
        "--ctx-size",
        str(settings.get("bundled_llama_ctx_size") or "4096"),
        "--parallel",
        "1",
    ]
    env = os.environ.copy()
    runtime_lib = str(Path(server).parent.parent / "lib")
    env["DYLD_LIBRARY_PATH"] = f"{runtime_lib}:{env.get('DYLD_LIBRARY_PATH', '')}".rstrip(":")
    process = subprocess.Popen(
        command,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
        text=True,
        start_new_session=True,
        env=env,
    )
    try:
        wait_for_llama_server(base_url, process)
        chunks = llm_text_chunks(text, int(settings.get("bundled_llama_chunk_chars") or 8500))
        results: list[dict[str, Any]] = []
        chunk_warnings: list[str] = []
        for index, chunk in enumerate(chunks, start=1):
            chunk_text = f"Document chunk {index} of {len(chunks)}. Extract only facts present in this chunk.\n\n{chunk}"
            try:
                results.append(
                    call_openai_compatible(
                        chunk_text,
                        {
                            "api_base_url": f"{base_url}/v1",
                            "api_model": model.stem,
                            "api_key": "",
                            "api_max_tokens": "512",
                        },
                        strict_schema=False,
                    )
                )
            except (OSError, RuntimeError, urllib.error.URLError, urllib.error.HTTPError, json.JSONDecodeError, KeyError) as exc:
                chunk_warnings.append(f"Bundled LLM chunk {index}/{len(chunks)} failed: {exc}")
        if not results:
            raise RuntimeError("; ".join(chunk_warnings) or "Bundled LLM returned no usable chunks.")
        merged = merge_llm_objects(results)
        merged.setdefault("warnings", [])
        merged["warnings"].extend(chunk_warnings)
        return merged
    finally:
        try:
            os.killpg(process.pid, 15)
        except OSError:
            process.terminate()
        try:
            process.wait(timeout=10)
        except subprocess.TimeoutExpired:
            try:
                os.killpg(process.pid, 9)
            except OSError:
                process.kill()
            try:
                process.wait(timeout=10)
            except subprocess.TimeoutExpired:
                pass


def cached_runtime_model(source: Path) -> Path:
    cache_dir = APP_SUPPORT / "models"
    cache_dir.mkdir(parents=True, exist_ok=True)
    target = cache_dir / source.name
    try:
        source_stat = source.stat()
        target_stat = target.stat() if target.exists() else None
        if target_stat and target_stat.st_size == source_stat.st_size and int(target_stat.st_mtime) >= int(source_stat.st_mtime):
            return target
        tmp = target.with_suffix(f"{target.suffix}.tmp")
        shutil.copy2(source, tmp)
        tmp.replace(target)
        return target
    except OSError:
        return source


def cached_runtime_tool(source: Path) -> Path:
    cache_root = APP_SUPPORT / "runtime-tools"
    bin_dir = cache_root / "bin"
    lib_dir = cache_root / "lib"
    bin_dir.mkdir(parents=True, exist_ok=True)
    target = bin_dir / source.name
    try:
        source_stat = source.stat()
        target_stat = target.stat() if target.exists() else None
        if not target_stat or target_stat.st_size != source_stat.st_size or int(target_stat.st_mtime) < int(source_stat.st_mtime):
            tmp = target.with_suffix(f"{target.suffix}.tmp")
            shutil.copy2(source, tmp)
            tmp.chmod(tmp.stat().st_mode | 0o755)
            tmp.replace(target)
        source_lib = source.parent.parent / "lib"
        if source_lib.exists():
            lib_dir.mkdir(parents=True, exist_ok=True)
            for item in source_lib.iterdir():
                if item.is_file() and item.suffix == ".dylib":
                    lib_target = lib_dir / item.name
                    item_stat = item.stat()
                    lib_stat = lib_target.stat() if lib_target.exists() else None
                    if not lib_stat or lib_stat.st_size != item_stat.st_size or int(lib_stat.st_mtime) < int(item_stat.st_mtime):
                        shutil.copy2(item, lib_target)
        return target
    except OSError:
        return source


def llm_extract(text: str, settings: dict[str, str]) -> tuple[dict[str, Any] | None, str | None]:
    provider = settings.get("provider") or "none"
    if provider == "none":
        return None, None
    try:
        if provider == "ollama":
            return call_ollama(text, settings), None
        if provider == "bundled_llama":
            return call_bundled_llama(text, settings), None
        if provider in {"openai", "openai_compatible"}:
            return call_openai_compatible(text, settings), None
    except urllib.error.HTTPError as exc:
        detail = ""
        try:
            detail = exc.read().decode("utf-8", errors="replace")[-1200:]
        except OSError:
            detail = ""
        return None, f"HTTP Error {exc.code}: {detail or exc.reason}"
    except (OSError, RuntimeError, urllib.error.URLError, json.JSONDecodeError, KeyError) as exc:
        return None, str(exc)
    return None, f"Unknown CV import provider: {provider}"


def normalize_llm_entry(entry: dict[str, Any]) -> dict[str, Any] | None:
    section_key = _text(entry.get("section_key"))
    if section_key not in SECTION_KEYS:
        section_key = "professional_positions"
    raw_text = _text(entry.get("raw_text") or entry.get("description") or entry.get("title"))
    title = _text(entry.get("title"))
    description = _text(entry.get("description"))
    if not raw_text and not title and not description:
        return None
    confidence = _text(entry.get("confidence")).lower()
    if confidence not in {"high", "medium", "low"}:
        confidence = "medium"
    return {
        "section_key": section_key,
        "subcategory": _text(entry.get("subcategory")) or None,
        "start_date": _text(entry.get("start_date")) or None,
        "end_date": _text(entry.get("end_date")) or None,
        "title": title or None,
        "organization": _text(entry.get("organization")) or None,
        "location": _text(entry.get("location")) or None,
        "role": _text(entry.get("role")) or None,
        "amount": _text(entry.get("amount")) or None,
        "description": description or None,
        "raw_text": raw_text or title or description,
        "confidence": confidence,
        "include_extended": 1,
        "include_long": 1,
        "include_short": 1 if section_key in {"education", "honors", "academic_appointments"} else 0,
        "include_biosketch": 1 if section_key == "honors" else 0,
        "language": _text(entry.get("language")) or "en",
        "source_note": "Imported from uploaded CV by LLM.",
    }


def merge_person(heuristic: dict[str, Any], llm: dict[str, Any] | None) -> dict[str, Any]:
    person = dict(heuristic)
    if llm:
        for field in PERSON_FIELDS:
            value = _text(llm.get(field))
            if value:
                person[field] = value
    for field in ("full_name", "display_name"):
        cleaned = clean_person_name(person.get(field))
        if cleaned:
            person[field] = cleaned
    for field in ("full_name", "display_name"):
        if _looks_like_document_title(_text(person.get(field))):
            person.pop(field, None)
    if not person.get("display_name") and person.get("full_name"):
        person["display_name"] = person["full_name"]
    return person


def insert_person(con: sqlite3.Connection, person: dict[str, Any]) -> None:
    values = {field: _text(person.get(field)) or None for field in PERSON_FIELDS}
    assignments = ", ".join(f"{field}=COALESCE(excluded.{field}, person.{field})" for field in PERSON_FIELDS)
    con.execute(
        f"""
        INSERT INTO person (id, {', '.join(PERSON_FIELDS)}, raw_json)
        VALUES (1, {', '.join('?' for _ in PERSON_FIELDS)}, ?)
        ON CONFLICT(id) DO UPDATE SET
        {assignments},
        raw_json=excluded.raw_json
        """,
        (*[values[field] for field in PERSON_FIELDS], json.dumps(values, ensure_ascii=False, indent=2)),
    )


def insert_entries(con: sqlite3.Connection, document_id: int, entries: list[dict[str, Any]]) -> int:
    count = 0
    seen: set[tuple[str, str]] = set()
    for entry in entries:
        normalized = normalize_llm_entry(entry)
        if not normalized:
            continue
        key = (str(normalized["section_key"]), str(normalized["raw_text"]).casefold())
        if key in seen:
            continue
        seen.add(key)
        existing = con.execute(
            "SELECT id FROM cv_entries WHERE section_key=? AND lower(raw_text)=? LIMIT 1",
            key,
        ).fetchone()
        if existing:
            continue
        con.execute(
            f"""
            INSERT INTO cv_entries (document_id, {', '.join(ENTRY_FIELDS)})
            VALUES (?, {', '.join('?' for _ in ENTRY_FIELDS)})
            """,
            (document_id, *[normalized.get(field) for field in ENTRY_FIELDS]),
        )
        count += 1
    return count


def normalize_publication(publication: dict[str, Any]) -> dict[str, Any] | None:
    publication = repair_publication_metadata(dict(publication))
    title = _clean_identifier(publication.get("title"))
    raw_citation = _text(publication.get("raw_citation"))
    authors = _clean_identifier(publication.get("authors"))
    venue = _clean_identifier(publication.get("venue"))
    year = _clean_identifier(publication.get("year"))
    if not title and not raw_citation:
        return None
    if not raw_citation:
        raw_citation = ". ".join(part for part in [authors, title, venue, year] if part)
    confidence = _text(publication.get("confidence")).lower()
    if confidence not in {"high", "medium", "low"}:
        confidence = "medium"
    category = _text(publication.get("category")).lower().replace(" ", "_").replace("-", "_") or "other"
    category_aliases = {
        "book_chapter": "books_chapters",
        "book_chapters": "books_chapters",
        "books_chapters": "books_chapters",
        "patent": "patents",
        "preprint": "preprints",
        "poster": "poster_presentations",
        "posters": "poster_presentations",
        "poster_presentation": "poster_presentations",
        "manuscript": "manuscripts_in_preparation",
        "manuscripts": "manuscripts_in_preparation",
        "manuscripts_under_review": "manuscripts_in_preparation",
    }
    category = category_aliases.get(category, category)
    if category not in {"peer_reviewed", "review", "abstract", "books_chapters", "patents", "preprints", "manuscripts_in_preparation", "poster_presentations", "other"}:
        category = "other"
    item_type = _text(publication.get("item_type")) or ("patent" if category == "patents" else "journal-article")
    return {
        "item_type": item_type,
        "category": category,
        "authors": authors,
        "title": title,
        "venue": venue,
        "year": year,
        "doi": _clean_identifier(publication.get("doi")),
        "pmid": _clean_identifier(publication.get("pmid")),
        "url": _text(publication.get("url")),
        "abstract": _text(publication.get("abstract")),
        "extra": _text(publication.get("extra")),
        "raw_citation": raw_citation,
        "confidence": confidence,
        "include_short": 0,
        "include_ultrashort": 0,
        "short_citation": _text(publication.get("short_citation")),
        "suppress_display": 0,
        "quality_note": _text(publication.get("quality_note")),
    }


def insert_publications(con: sqlite3.Connection, document_id: int, publications: list[dict[str, Any]]) -> int:
    count = 0
    seen: set[tuple[str, str, str]] = set()
    for publication in publications:
        normalized = normalize_publication(publication)
        if not normalized:
            continue
        key = (
            normalized["doi"].casefold(),
            normalized["pmid"].casefold(),
            normalized["raw_citation"].casefold(),
        )
        if key in seen:
            continue
        seen.add(key)
        if normalized["doi"]:
            existing = con.execute(
                "SELECT id FROM publications WHERE lower(COALESCE(doi, '')) = ? LIMIT 1",
                (normalized["doi"].casefold(),),
            ).fetchone()
            if existing:
                continue
        if normalized["pmid"]:
            existing = con.execute("SELECT id FROM publications WHERE pmid = ? LIMIT 1", (normalized["pmid"],)).fetchone()
            if existing:
                continue
        existing = con.execute(
            "SELECT id FROM publications WHERE lower(raw_citation) = ? LIMIT 1",
            (normalized["raw_citation"].casefold(),),
        ).fetchone()
        if existing:
            continue
        con.execute(
            f"""
            INSERT INTO publications (document_id, source, {', '.join(PUBLICATION_FIELDS)})
            VALUES (?, 'document_llm', {', '.join('?' for _ in PUBLICATION_FIELDS)})
            """,
            (document_id, *[normalized[field] for field in PUBLICATION_FIELDS]),
        )
        count += 1
    return count


def insert_narrative_report(con: sqlite3.Connection, report: dict[str, Any] | None) -> int:
    if not isinstance(report, dict):
        return 0
    body = _text(report.get("body"))
    body_de = _text(report.get("body_de"))
    if not body and not body_de:
        return 0
    title = _text(report.get("title")) or "Narrative Report"
    title_de = _text(report.get("title_de")) or "Freie Stellungname"
    con.execute(
        """
        INSERT INTO narrative_reports (id, title, body, title_de, body_de, updated_at)
        VALUES (1, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        ON CONFLICT(id) DO UPDATE SET
          title=excluded.title,
          body=excluded.body,
          title_de=excluded.title_de,
          body_de=excluded.body_de,
          updated_at=CURRENT_TIMESTAMP
        """,
        (title, body, title_de, body_de),
    )
    return 1


def find_publication_id(con: sqlite3.Connection, citation: dict[str, Any]) -> int | None:
    doi = _clean_identifier(citation.get("doi")).casefold()
    pmid = _clean_identifier(citation.get("pmid"))
    raw_citation = _text(citation.get("raw_citation")).casefold()
    if doi:
        row = con.execute(
            "SELECT id FROM publications WHERE lower(COALESCE(doi, '')) = ? ORDER BY id LIMIT 1",
            (doi,),
        ).fetchone()
        if row:
            return int(row["id"])
    if pmid:
        row = con.execute("SELECT id FROM publications WHERE pmid = ? ORDER BY id LIMIT 1", (pmid,)).fetchone()
        if row:
            return int(row["id"])
    if raw_citation:
        row = con.execute(
            "SELECT id FROM publications WHERE lower(raw_citation) = ? ORDER BY id LIMIT 1",
            (raw_citation,),
        ).fetchone()
        if row:
            return int(row["id"])
    return None


def normalize_contribution(contribution: dict[str, Any]) -> dict[str, str] | None:
    title = _text(contribution.get("title"))
    narrative = _text(contribution.get("narrative"))
    raw_text = _text(contribution.get("raw_text"))
    if not title and raw_text:
        title = raw_text[:120]
    if not title or not (narrative or raw_text):
        return None
    return {
        "title": title,
        "narrative": narrative or raw_text,
        "citations_json": json.dumps(contribution.get("citations") or [], ensure_ascii=False),
    }


def insert_contributions(con: sqlite3.Connection, document_id: int, contributions: list[dict[str, Any]]) -> int:
    count = 0
    seen: set[tuple[str, str]] = set()
    for index, contribution in enumerate(contributions, start=1):
        normalized = normalize_contribution(contribution)
        if not normalized:
            continue
        key = (normalized["title"].casefold(), normalized["narrative"].casefold())
        if key in seen:
            continue
        seen.add(key)
        existing = con.execute(
            "SELECT id FROM biosketch_contributions WHERE lower(title)=? AND lower(narrative)=? LIMIT 1",
            key,
        ).fetchone()
        if existing:
            continue
        cursor = con.execute(
            """
            INSERT INTO biosketch_contributions (document_id, ordinal, title, narrative, citations_json)
            VALUES (?, ?, ?, ?, ?)
            """,
            (
                document_id,
                contribution.get("ordinal") or index,
                normalized["title"],
                normalized["narrative"],
                normalized["citations_json"],
            ),
        )
        contribution_id = int(cursor.lastrowid)
        for citation_index, citation in enumerate(contribution.get("citations") or [], start=1):
            if not isinstance(citation, dict):
                continue
            raw_citation = _text(citation.get("raw_citation"))
            if not raw_citation:
                continue
            label = _text(citation.get("label")) or str(citation_index)
            publication_id = find_publication_id(con, citation)
            con.execute(
                """
                INSERT OR IGNORE INTO biosketch_contribution_publications
                  (contribution_id, citation_label, publication_id, raw_citation, pmid, doi)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (
                    contribution_id,
                    label,
                    publication_id,
                    raw_citation,
                    _clean_identifier(citation.get("pmid")),
                    _clean_identifier(citation.get("doi")),
                ),
            )
        count += 1
    return count


def import_cv_file(con: sqlite3.Connection, path: Path, original_filename: str, settings: dict[str, str]) -> dict[str, Any]:
    text, source_format = extract_text(path)
    if not text.strip():
        raise RuntimeError("No text could be extracted from this CV.")

    imported_at = dt.datetime.now(dt.timezone.utc).isoformat()
    slug_base = re.sub(r"[^a-z0-9]+", "_", Path(original_filename).stem.casefold()).strip("_") or "uploaded_cv"
    slug = f"uploaded_cv_{slug_base}_{dt.datetime.now().strftime('%Y%m%d%H%M%S')}"
    cursor = con.execute(
        """
        INSERT INTO documents (slug, title, source_path, source_format, imported_at, notes)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (slug, original_filename, str(path), source_format, imported_at, "Imported from user-uploaded CV."),
    )
    document_id = int(cursor.lastrowid)

    heuristic = heuristic_person(text)
    heuristic_rows, skipped_sections = heuristic_entries(text)
    heuristic_pubs = heuristic_publications(text)
    heuristic_biosketch = heuristic_contributions(text)
    heuristic_report = heuristic_narrative_report(text)
    llm_data, llm_warning = llm_extract(text, settings)
    llm_rows = [normalize_llm_entry(row) for row in (llm_data or {}).get("entries", []) if isinstance(row, dict)]
    llm_entries = [row for row in llm_rows if row]
    llm_contributions = [
        contribution
        for contribution in (llm_data or {}).get("contributions", [])
        if isinstance(contribution, dict)
    ]
    llm_publications = [
        publication
        for publication in (llm_data or {}).get("publications", [])
        if isinstance(publication, dict)
    ]
    entries = llm_entries if llm_entries else heuristic_rows
    contributions = llm_contributions if llm_contributions else heuristic_biosketch
    person = merge_person(heuristic, (llm_data or {}).get("person") if isinstance(llm_data, dict) else None)
    llm_report = (llm_data or {}).get("narrative_report") if isinstance(llm_data, dict) else None
    report = llm_report if isinstance(llm_report, dict) and (_text(llm_report.get("body")) or _text(llm_report.get("body_de"))) else heuristic_report

    insert_person(con, person)
    inserted = insert_entries(con, document_id, entries)
    publications = llm_publications if llm_publications else heuristic_pubs
    if not publications and contributions:
        publications = contribution_publications(contributions)
    publications_inserted = insert_publications(con, document_id, publications)
    contributions_inserted = insert_contributions(con, document_id, contributions)
    narrative_imported = insert_narrative_report(con, report)

    warnings: list[str] = []
    if llm_warning:
        warnings.append(f"LLM import fell back to heuristic parsing: {llm_warning}")
    if not llm_entries and settings.get("provider") != "none":
        warnings.append("No usable structured entries were returned by the LLM; heuristic entries were imported.")
    if not llm_publications and heuristic_pubs:
        warnings.append(f"Imported {len(heuristic_pubs)} publication-looking citations with the heuristic parser.")
    skipped_publications = int((llm_data or {}).get("skipped_publication_count") or 0) if isinstance(llm_data, dict) else 0
    skipped_publications += skipped_sections
    if skipped_publications:
        warnings.append(f"Some publication-looking sections/items were not parsed ({skipped_publications}).")
    llm_warnings = (llm_data or {}).get("warnings", []) if isinstance(llm_data, dict) else []
    if not isinstance(llm_warnings, list):
        llm_warnings = []
    for warning in warnings + [str(item) for item in llm_warnings if item]:
        con.execute(
            """
            INSERT INTO import_warnings (document_id, warning_type, message, raw_text)
            VALUES (?, 'cv_import', ?, ?)
            """,
            (document_id, warning[:1000], None),
        )

    return {
        "document_id": document_id,
        "entries_inserted": inserted,
        "contributions_inserted": contributions_inserted,
        "publications_inserted": publications_inserted,
        "narrative_imported": narrative_imported,
        "person_fields": len([value for value in person.values() if value]),
        "provider": settings.get("provider") or "none",
        "used_llm": bool(llm_data) and not llm_warning,
        "warnings": warnings,
    }

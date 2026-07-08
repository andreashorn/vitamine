#!/usr/bin/env python3
"""Build a compact short CV from entries and publications flagged for short export."""

from __future__ import annotations

import html
import json
import re
import sqlite3
import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from vitamine.scripts.export_utils import compile_typst_if_available
from vitamine.paths import OUTPUT, ROOT, active_db_path, output_ref

DB = active_db_path()
LANG = "en"


SECTION_TITLES = {
    "education": "Education and Training",
    "postdoctoral_training": "Postdoctoral Training",
    "academic_appointments": "Positions and Appointments",
    "hospital_appointments": "Hospital Appointments",
    "honors": "Selected Honors",
    "funding": "Selected Funding",
    "mentoring": "Mentoring",
    "invited_presentations": "Selected Presentations",
}

SECTION_TITLES_DE = {
    "education": "Ausbildung",
    "postdoctoral_training": "Postdoktorale Ausbildung",
    "academic_appointments": "Positionen und Berufungen",
    "hospital_appointments": "Klinische Positionen",
    "honors": "Ausgewählte Auszeichnungen",
    "funding": "Ausgewählte Forschungsförderung",
    "mentoring": "Betreuung",
    "invited_presentations": "Ausgewählte Vorträge",
}

SECTION_ORDER = [
    "education",
    "postdoctoral_training",
    "academic_appointments",
    "hospital_appointments",
    "funding",
    "honors",
    "mentoring",
    "invited_presentations",
]


def connect() -> sqlite3.Connection:
    con = sqlite3.connect(DB)
    con.row_factory = sqlite3.Row
    ensure_export_columns(con)
    return con


def ensure_export_columns(con: sqlite3.Connection) -> None:
    existing = {row[1] for row in con.execute("PRAGMA table_info(publications)").fetchall()}
    columns = {
        "short_selected_order": "INTEGER",
        "ultrashort_selected_order": "INTEGER",
    }
    for column, definition in columns.items():
        if column not in existing:
            con.execute(f"ALTER TABLE publications ADD COLUMN {column} {definition}")
    con.execute(
        """
        CREATE TABLE IF NOT EXISTS export_settings (
          profile TEXT PRIMARY KEY,
          publication_limit INTEGER NOT NULL DEFAULT 10,
          authorship_filter TEXT NOT NULL DEFAULT 'first_last'
        )
        """
    )
    con.execute(
        """
        INSERT OR IGNORE INTO export_settings (profile, publication_limit, authorship_filter)
        VALUES ('short', 10, 'first_last')
        """
    )


def clean(value: str | None) -> str:
    return (value or "").strip()


def row_value(row: sqlite3.Row, field: str) -> str:
    if LANG == "de":
        german = f"{field}_de"
        if german in row.keys() and clean(row[german]):
            return clean(row[german])
    return clean(row[field])


def section_title(section_key: str) -> str:
    labels = SECTION_TITLES_DE if LANG == "de" else SECTION_TITLES
    return labels.get(section_key, section_key.replace("_", " ").title())


def typ(value: str | None) -> str:
    return json.dumps(clean(value), ensure_ascii=False)


def text(value: str | None, *, bold: bool = False, size: str | None = None) -> str:
    args = []
    if bold:
        args.append('weight: "bold"')
    if size:
        args.append(f"size: {size}")
    args.append(typ(value))
    return f"#text({', '.join(args)})"


def citation_cell(value: str | None) -> str:
    value = clean(value)
    value = re.sub(r"\s+([,.;:])", r"\1", value)
    value = re.sub(r",(?=\S)", ", ", value)
    value = re.sub(r"\s{2,}", " ", value)
    return value.rstrip(" .")


def sentence_part(value: str) -> str:
    value = citation_cell(value)
    return value if not value or value.endswith((".", "?", "!")) else f"{value}."


def impact_factor_label(row: sqlite3.Row) -> str:
    value = row["impact_factor"] if "impact_factor" in row.keys() else None
    if value in (None, ""):
        return ""
    try:
        formatted = f"{float(value):g}"
    except (TypeError, ValueError):
        formatted = clean(str(value))
    year = citation_cell(row["impact_factor_year"] if "impact_factor_year" in row.keys() else "")
    return f"IF {formatted} ({year})" if year else f"IF {formatted}"


def typst_rich_text(value: str | None, *, bold_names: bool = False, underline: bool = False, italic: bool = False, size: str = "9pt") -> str:
    value = clean(value)
    if not value:
        return text("", size=size)
    if bold_names:
        pieces = []
        cursor = 0
        for match in re.finditer(r"\b(?:Andreas\s+Horn|Horn\s+A\.?|Horn)\b", value):
            if match.start() > cursor:
                before = value[cursor : match.start()]
                stripped = before.rstrip()
                if stripped:
                    pieces.append(text(stripped, size=size))
                if before and before[-1].isspace():
                    pieces.append("#h(0.28em)")
            pieces.append(text(match.group(0), bold=True, size=size))
            cursor = match.end()
        if cursor < len(value):
            after = value[cursor:]
            if after and after[0].isspace():
                pieces.append("#h(0.28em)")
                after = after.lstrip()
            if after:
                pieces.append(text(after, size=size))
        return "".join(pieces)
    body = text(value, size=size)
    if italic:
        body = f"#emph[{body}]"
    if underline:
        body = f"#underline[{body}]"
    return body


def period(row: sqlite3.Row) -> str:
    start = clean(row["start_date"])
    end = clean(row["end_date"])
    if start and end:
        return f"{start}-{end}"
    return start


def year_label(row: sqlite3.Row) -> str:
    start = clean(row["start_date"])
    end = clean(row["end_date"])
    if start and end and start != end:
        return f"{start}-{end}"
    return start or end


def normalize_text(value: str) -> str:
    return re.sub(r"\W+", " ", value.casefold()).strip()


def entry_detail(row: sqlite3.Row) -> str:
    title = row_value(row, "title")
    organization = row_value(row, "organization")
    role = row_value(row, "role")
    amount = row_value(row, "amount")
    description = row_value(row, "description")
    if row["section_key"] == "honors":
        pieces = [title]
        if organization and normalize_text(organization) not in normalize_text(title):
            pieces.append(organization)
        return ", ".join(piece for piece in pieces if piece)
    if "|" in description:
        desc_parts = [clean(part) for part in description.split("|") if clean(part)]
        if desc_parts:
            title = title or desc_parts[0]
            if len(desc_parts) >= 2 and not role:
                role = desc_parts[1]
            if len(desc_parts) >= 3 and not organization:
                organization = desc_parts[2]
    parts = [title, organization, role, amount]
    if description and "|" not in description:
        parts.append(description)
    return "; ".join(part for part in parts if part)


def citation(row: sqlite3.Row) -> str:
    authors = citation_cell(row["authors"])
    if authors:
        author_parts = [part.strip() for part in authors.split(",") if part.strip()]
        if len(author_parts) > 4:
            authors = ", ".join(author_parts[:3]) + ", et al."
    parts = [authors, citation_cell(row["title"]), citation_cell(row["venue"]), citation_cell(row["year"]), impact_factor_label(row)]
    out = ". ".join(part for part in parts if part)
    if citation_cell(row["doi"]):
        out = f"{out}. doi:{citation_cell(row['doi'])}"
    return out


def typst_publication_citation(row: sqlite3.Row) -> str:
    parts = []
    authors = citation_cell(row["authors"])
    if authors:
        parts.append(typst_rich_text(sentence_part(authors), bold_names=True))
    title = citation_cell(row["title"])
    if title:
        parts.append(text(sentence_part(title), size="9pt"))
    venue = citation_cell(row["venue"])
    if venue:
        venue = venue.title() if venue.isupper() else venue
        parts.append(typst_rich_text(sentence_part(venue), italic=True, underline=True))
    year = citation_cell(row["year"])
    if year:
        parts.append(text(sentence_part(year), size="9pt"))
    impact = impact_factor_label(row)
    if impact:
        parts.append(text(sentence_part(impact), size="9pt"))
    doi = citation_cell(row["doi"])
    if doi:
        parts.append(text(sentence_part(f"doi:{doi}"), size="9pt"))
    return "#h(0.28em)".join(parts)


def load_data() -> tuple[sqlite3.Row | None, list[sqlite3.Row], list[sqlite3.Row]]:
    with connect() as con:
        person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
        settings = con.execute("SELECT publication_limit FROM export_settings WHERE profile='short'").fetchone()
        publication_limit = int(settings["publication_limit"] if settings else 10)
        entries = con.execute(
            """
            SELECT *
            FROM cv_entries
            WHERE include_short=1
            ORDER BY section_key, start_date, id
            """
        ).fetchall()
        pubs = con.execute(
            """
            SELECT *
            FROM publications
            WHERE include_short=1
              AND COALESCE(suppress_display, 0)=0
              AND category='peer_reviewed'
            ORDER BY COALESCE(short_selected_order, selected_order, 999), CAST(year AS INTEGER) DESC, id DESC
            LIMIT ?
            """
            ,
            (publication_limit,),
        ).fetchall()
    return person, entries, pubs


def grouped_sections(entries: list[sqlite3.Row]) -> list[tuple[str, list[sqlite3.Row]]]:
    by_section: dict[str, list[sqlite3.Row]] = {}
    seen: set[tuple[str, str]] = set()
    for row in entries:
        detail = entry_detail(row)
        key = (row["section_key"], normalize_text(f"{year_label(row)} {detail}"))
        if key in seen:
            continue
        seen.add(key)
        by_section.setdefault(row["section_key"], []).append(row)
    ordered = []
    for section_key in SECTION_ORDER:
        if section_key in by_section:
            rows = by_section.pop(section_key)
            if section_key == "honors":
                rows = sorted(rows, key=lambda row: (year_label(row), row["id"]), reverse=True)
            ordered.append((section_key, rows))
    ordered.extend((key, rows) for key, rows in sorted(by_section.items()))
    return ordered


def build_html() -> str:
    person, entries, pubs = load_data()
    name = clean(person["display_name"] if person else "") or "Andreas Horn"
    title = clean(person["position_title"] if person else "")
    body = [f"<h1>{html.escape(name)}</h1>"]
    if title:
        body.append(f"<p class=\"subtitle\">{html.escape(title)}</p>")
    for section_key, rows in grouped_sections(entries):
        body.append(f"<h2>{html.escape(section_title(section_key))}</h2>")
        body.append("<table>")
        for row in rows:
            body.append(f"<tr><td>{html.escape(year_label(row))}</td><td>{html.escape(entry_detail(row))}</td></tr>")
        body.append("</table>")
    if pubs:
        body.append(f"<h2>{html.escape('Ausgewählte Publikationen' if LANG == 'de' else 'Selected Publications')}</h2><ol>")
        for row in pubs:
            body.append(f"<li>{html.escape(citation(row))}</li>")
        body.append("</ol>")
    return f"""<!doctype html>
<html lang="{LANG}">
<head>
  <meta charset="utf-8">
  <title>Short CV</title>
  <style>
    body {{ font-family: Arial, Helvetica, sans-serif; max-width: 900px; margin: 28px auto; color: #111; font-size: 12px; line-height: 1.32; }}
    h1 {{ font-size: 22px; margin: 0 0 4px; }}
    .subtitle {{ font-weight: 700; margin-bottom: 16px; }}
    h2 {{ font-size: 14px; margin: 16px 0 6px; border-bottom: 1px solid #999; }}
    table {{ width: 100%; border-collapse: collapse; }}
    td {{ padding: 3px 6px; vertical-align: top; border-bottom: 1px solid #ddd; }}
    td:first-child {{ width: 150px; font-weight: 700; }}
  </style>
</head>
<body>
{''.join(body)}
</body>
</html>
"""


def build_typst() -> str:
    person, entries, pubs = load_data()
    name = clean(person["display_name"] if person else "") or "Andreas Horn"
    title = clean(person["position_title"] if person else "")
    lines = [
        '#set page(width: 8.5in, height: 11in, margin: (left: 0.65in, right: 0.65in, top: 0.65in, bottom: 0.58in))',
        f'#set text(font: "Arial", size: 10pt, lang: "{LANG}")',
        "#set par(leading: 0.45em)",
        text(name, bold=True, size="16pt"),
    ]
    if title:
        lines.append(text(title, bold=True))
    for section_key, rows in grouped_sections(entries):
        lines.append(f"\n#line(length: 100%)\n{text(section_title(section_key), bold=True)}")
        cells = []
        for row in rows:
            cells.append(f"[{text(year_label(row), bold=True)}]")
            cells.append(f"[{text(entry_detail(row))}]")
        lines.append("#grid(columns: (1.45in, 5.55in), gutter: 0.2in, row-gutter: 0.045in,\n" + ",\n".join(cells) + "\n)")
    if pubs:
        lines.append(f"\n#line(length: 100%)\n{text('Ausgewählte Publikationen' if LANG == 'de' else 'Selected Publications', bold=True)}")
        for index, row in enumerate(pubs, 1):
            lines.append("#grid(columns: (0.3in, 6.8in), gutter: 0.08in,\n"
                         f"  [{text(str(index)+'.')}],\n  [{typst_publication_citation(row)}]\n)")
    return "\n".join(lines) + "\n"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D7DCE2") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_font(paragraph, *, size: float = 8.7, bold: bool = False, color: str = "111827") -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.03


def set_run_font(run, *, size: float = 8.0, bold: bool = False, italic: bool = False, underline: bool = False, color: str = "111827") -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = RGBColor.from_string(color)


def add_docx_piece(paragraph, value: str, *, bold_names: bool = False, italic: bool = False, underline: bool = False) -> None:
    if not value:
        return
    if not bold_names:
        set_run_font(paragraph.add_run(value), italic=italic, underline=underline)
        return
    cursor = 0
    for match in re.finditer(r"\b(?:Andreas\s+Horn|Horn\s+A\.?|Horn)\b", value):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(value[cursor : match.start()]), italic=italic, underline=underline)
        set_run_font(paragraph.add_run(match.group(0)), bold=True, italic=italic, underline=underline)
        cursor = match.end()
    if cursor < len(value):
        set_run_font(paragraph.add_run(value[cursor:]), italic=italic, underline=underline)


def add_publication_docx(paragraph, row: sqlite3.Row) -> None:
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.03
    parts: list[tuple[str, bool, bool, bool]] = []
    authors = citation_cell(row["authors"])
    if authors:
        parts.append((sentence_part(authors), True, False, False))
    title = citation_cell(row["title"])
    if title:
        parts.append((sentence_part(title), False, False, False))
    venue = citation_cell(row["venue"])
    if venue:
        venue = venue.title() if venue.isupper() else venue
        parts.append((sentence_part(venue), False, True, True))
    year = citation_cell(row["year"])
    if year:
        parts.append((sentence_part(year), False, False, False))
    impact = impact_factor_label(row)
    if impact:
        parts.append((sentence_part(impact), False, False, False))
    doi = citation_cell(row["doi"])
    if doi:
        parts.append((sentence_part(f"doi:{doi}"), False, False, False))
    for index, (value, bold_names, italic, underline) in enumerate(parts):
        if index:
            set_run_font(paragraph.add_run(" "))
        add_docx_piece(paragraph, value, bold_names=bold_names, italic=italic, underline=underline)


def add_compact_heading(doc: Document, label: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(9.3)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("111827")


def set_table_width(table, widths) -> None:
    table.autofit = False
    total = sum(int(width) for width in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width)))
            tc_w.set(qn("w:type"), "dxa")


def add_entries_table(doc: Document, rows: list[sqlite3.Row]) -> None:
    widths = [Inches(1.18), Inches(6.12)]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, widths)
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = year_label(row)
        cells[1].text = entry_detail(row)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if index % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
        for paragraph in cells[0].paragraphs:
            set_paragraph_font(paragraph, size=8.3, bold=True, color="334155")
        for paragraph in cells[1].paragraphs:
            set_paragraph_font(paragraph, size=8.4)
    set_table_width(table, widths)


def add_publications_table(doc: Document, pubs: list[sqlite3.Row]) -> None:
    widths = [Inches(0.32), Inches(6.98)]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, widths)
    for index, row in enumerate(pubs, 1):
        cells = table.add_row().cells
        cells[0].text = f"{index}."
        cells[1].text = ""
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, top=60, start=75, bottom=60, end=75)
        for paragraph in cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_paragraph_font(paragraph, size=8.0, bold=True, color="334155")
        for paragraph in cells[1].paragraphs:
            add_publication_docx(paragraph, row)
    set_table_width(table, widths)


def build_docx(path: Path) -> Path:
    person, entries, pubs = load_data()
    name = clean(person["display_name"] if person else "") or "Andreas Horn"
    title = clean(person["position_title"] if person else "")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(8.6)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(1)
    run = title_p.add_run(name)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(15)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("111827")
    if title:
        subtitle = doc.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(6)
        run = subtitle.add_run(title)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(8.8)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("334155")

    for section_key, rows in grouped_sections(entries):
        add_compact_heading(doc, section_title(section_key))
        add_entries_table(doc, rows)
    if pubs:
        add_compact_heading(doc, "Ausgewählte Publikationen" if LANG == "de" else "Selected Publications")
        add_publications_table(doc, pubs)

    doc.core_properties.title = "Short CV"
    doc.core_properties.author = name
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def output_stem() -> str:
    return "short_cv_de" if LANG == "de" else "short_cv"


def build(lang: str = "en") -> dict[str, str]:
    global LANG
    LANG = "de" if lang == "de" else "en"
    OUTPUT.mkdir(parents=True, exist_ok=True)
    stem = output_stem()
    html_path = OUTPUT / f"{stem}.html"
    typ_path = OUTPUT / f"{stem}.typ"
    pdf_path = OUTPUT / f"{stem}.pdf"
    docx_path = OUTPUT / f"{stem}.docx"
    html_path.write_text(build_html(), encoding="utf-8")
    typ_path.write_text(build_typst(), encoding="utf-8")
    pdf, warning = compile_typst_if_available(typ_path, pdf_path, ROOT)
    docx = build_docx(docx_path)
    result = {
        "html": f"output/{output_ref(html_path)}",
        "typst": f"output/{output_ref(typ_path)}",
    }
    if pdf:
        result["pdf"] = f"output/{output_ref(pdf)}"
    if docx:
        result["docx"] = f"output/{output_ref(docx)}"
    warnings = [item for item in (warning,) if item]
    if warnings:
        result["warning"] = " ".join(warnings)
    return result


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--lang", choices=["en", "de"], default="en")
    args = parser.parse_args()
    for name, path in build(args.lang).items():
        print(f"{name}: {path}")

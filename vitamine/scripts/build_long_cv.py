#!/usr/bin/env python3
"""Build the long CV from the SQLite database."""

from __future__ import annotations

import datetime as dt
import argparse
import html
import json
import re
import sqlite3
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from vitamine.scripts.export_utils import compile_typst_if_available, markdown_to_html_body
from vitamine.paths import OUTPUT, ROOT, active_db_path, output_ref

DB = active_db_path()

LANG = "en"


SECTION_ORDER = [
    ("education", "Education"),
    ("postdoctoral_training", "Postdoctoral Training"),
    ("academic_appointments", "Faculty Academic Appointments"),
    ("hospital_appointments", "Appointments at Hospitals/Affiliated Institutions"),
    ("professional_positions", "Other Professional Positions"),
    ("committee_service", "Committee Service"),
    ("professional_societies", "Professional Societies"),
    ("grant_review", "Grant Review Activities"),
    ("editorial_activities", "Editorial Activities"),
    ("honors", "Honors and Prizes"),
    ("funding", "Research Funding"),
    ("teaching", "Teaching of Students in Courses"),
    ("mentoring", "Research Supervisory and Training Responsibilities"),
    ("invited_presentations", "Report of Regional, National and International Invited Teaching and Presentations"),
    ("clinical_activities", "Report of Clinical Activities and Innovations"),
    ("education_innovations", "Report of Teaching and Education Innovations"),
    ("community_service", "Report of Education of Patients and Service to the Community"),
]

SECTION_LABELS_DE = {
    "education": "Ausbildung",
    "postdoctoral_training": "Postdoktorale Ausbildung",
    "academic_appointments": "Akademische Berufungen",
    "hospital_appointments": "Positionen an Kliniken und affiliierten Institutionen",
    "professional_positions": "Weitere berufliche Positionen",
    "committee_service": "Gremienarbeit",
    "professional_societies": "Fachgesellschaften",
    "grant_review": "Gutachtertätigkeiten für Forschungsförderung",
    "editorial_activities": "Editoriale Tätigkeiten",
    "honors": "Auszeichnungen und Preise",
    "funding": "Forschungsförderung",
    "teaching": "Lehre in Kursen",
    "mentoring": "Betreuung und Ausbildung",
    "invited_presentations": "Bericht über regionale, nationale und internationale eingeladene Lehre und Vorträge",
    "clinical_activities": "Bericht über klinische Tätigkeiten und Innovationen",
    "education_innovations": "Bericht über Lehr- und Ausbildungsinnovationen",
    "community_service": "Bericht über Patientenaufklärung und gesellschaftliches Engagement",
}

TEXT = {
    "en": {
        "faculty": "The Faculty of Medicine of University Cologne",
        "cv": "Curriculum Vitae",
        "date_prepared": "Date Prepared:",
        "name": "Name:",
        "office_address": "Office Address:",
        "home_address": "Home Address:",
        "work_phone": "Work Phone:",
        "work_email": "Work Email:",
        "place_of_birth": "Place of Birth:",
        "scholarship": "Report of Scholarship",
        "peer_reviewed": "Peer-Reviewed Scholarship in print or other media:",
        "other_scholarship": "Other Scholarship",
        "no_sponsor": "No presentations below were sponsored by 3rd parties/outside entities.",
        "achievements": "Achievements",
        "ad_hoc_reviewer": "Ad hoc Reviewer",
        "other_editorial_roles": "Other Editorial Roles",
        "other_trainees": "Other Formally Supervised Trainees",
        "narrative_report": "Narrative Report",
        "dates": "Dates",
        "role_title": "Role / Title",
        "field_details": "Field / Details",
        "institution": "Institution / Organization",
        "funding_source": "Funding Source",
        "role": "Role",
        "amount": "Amount",
        "details": "Details",
    },
    "de": {
        "faculty": "Medizinische Fakultät der Universität zu Köln",
        "cv": "Lebenslauf",
        "date_prepared": "Erstellt am:",
        "name": "Name:",
        "office_address": "Dienstadresse:",
        "home_address": "Privatadresse:",
        "work_phone": "Telefon dienstlich:",
        "work_email": "E-Mail dienstlich:",
        "place_of_birth": "Geburtsort:",
        "scholarship": "Publikationsbericht",
        "peer_reviewed": "Begutachtete wissenschaftliche Publikationen:",
        "other_scholarship": "Weitere wissenschaftliche Beiträge",
        "no_sponsor": "Die unten aufgeführten Vorträge wurden nicht durch Dritte/externe Einrichtungen gesponsert.",
        "achievements": "Erfolge",
        "ad_hoc_reviewer": "Ad-hoc-Gutachter",
        "other_editorial_roles": "Weitere editoriale Tätigkeiten",
        "other_trainees": "Weitere formal betreute Personen",
        "narrative_report": "Narrativer Bericht",
        "dates": "Zeitraum",
        "role_title": "Rolle / Titel",
        "field_details": "Fach / Details",
        "institution": "Institution / Organisation",
        "funding_source": "Fördermittelgeber",
        "role": "Rolle",
        "amount": "Betrag",
        "details": "Details",
    },
}

THREE_COLUMN_SECTIONS = {
    "committee_service",
    "professional_societies",
    "grant_review",
    "teaching",
}


def clean(value: str | None) -> str:
    return value or ""


def clean_cell(value: str | None) -> str:
    text = clean(value).strip()
    if text.startswith(">"):
        text = text[1:].strip()
    text = text.replace("^st^", "st").replace("^nd^", "nd").replace("^rd^", "rd").replace("^th^", "th")
    return text


def tr(key: str) -> str:
    return TEXT.get(LANG, TEXT["en"])[key]


def localized_section_title(section_key: str, english_title: str) -> str:
    if LANG == "de":
        return SECTION_LABELS_DE.get(section_key, english_title)
    return english_title


def row_value(row: sqlite3.Row, field: str) -> str:
    if LANG == "de":
        german = f"{field}_de"
        if german in row.keys() and clean_cell(row[german]):
            return clean_cell(row[german])
    return clean_cell(row[field] if field in row.keys() else "")


def period(start: str | None, end: str | None) -> str:
    if start and end:
        return f"{start}-{end}"
    if start:
        return f"{start}-"
    return ""


def sort_date(value: str | None) -> tuple[int, int, int]:
    if not value:
        return (9999, 12, 31)
    value = value.strip().rstrip(",")
    parts = value.split("/")
    if len(parts) == 2 and all(part.isdigit() for part in parts):
        month, year = [int(part) for part in parts]
        return (year, month, 1)
    if len(parts) == 3 and all(part.isdigit() for part in parts):
        month, day, year = [int(part) for part in parts]
        if year < 100:
            year += 2000 if year < 40 else 1900
        return (year, month, day)
    if value[:4].isdigit():
        return (int(value[:4]), 1, 1)
    return (9999, 12, 31)


def md_escape(value: str | None) -> str:
    text = clean(value)
    return text.replace("|", "\\|")


def html_escape(value: str | None) -> str:
    return html.escape(clean(value))


def typ_string(value: str | None) -> str:
    text = clean(value)
    text = text.replace("<br>", "\n")
    text = text.replace("\\|", "|")
    return json.dumps(text, ensure_ascii=False)


def typ_text(value: str | None, *, bold: bool = False, size: str | None = None) -> str:
    args = []
    if bold:
        args.append('weight: "bold"')
    if size:
        args.append(f"size: {size}")
    args.append(typ_string(value))
    return f"#text({', '.join(args)})"


def connect() -> sqlite3.Connection:
    con = sqlite3.connect(DB)
    con.row_factory = sqlite3.Row
    con.execute(
        """
        CREATE TABLE IF NOT EXISTS narrative_reports (
          id INTEGER PRIMARY KEY CHECK (id = 1),
          title TEXT NOT NULL DEFAULT 'Narrative Report',
          body TEXT NOT NULL DEFAULT '',
          title_de TEXT,
          body_de TEXT,
          updated_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    existing = {row[1] for row in con.execute("PRAGMA table_info(narrative_reports)").fetchall()}
    if "title_de" not in existing:
        con.execute("ALTER TABLE narrative_reports ADD COLUMN title_de TEXT")
    if "body_de" not in existing:
        con.execute("ALTER TABLE narrative_reports ADD COLUMN body_de TEXT")
    return con


def person_block(con: sqlite3.Connection) -> list[str]:
    person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
    today = dt.datetime.now().strftime("%d.%m.%Y") if LANG == "de" else dt.datetime.now().strftime("%B %-d, %Y")
    if not person:
        return [f"**{tr('faculty')}**", "", f"# {tr('cv')}"]
    rows = [
        (tr("date_prepared").rstrip(":"), today),
        (tr("name").rstrip(":"), person["display_name"] or person["full_name"]),
        (tr("office_address").rstrip(":"), person["office_address"]),
        (tr("home_address").rstrip(":"), person["home_address"]),
        (tr("work_phone").rstrip(":"), person["work_phone"]),
        (tr("work_email").rstrip(":"), person["work_email"]),
        (tr("place_of_birth").rstrip(":"), person["place_of_birth"]),
    ]
    out = [f"<div class=\"cv-kicker\">{tr('faculty')}</div>", "", f"# {tr('cv')}", ""]
    out.append("|  |  |")
    out.append("| --- | --- |")
    for label, value in rows:
        if value:
            out.append(f"| **{label}:** | {md_escape(value)} |")
    return out


def entry_rows(con: sqlite3.Connection, section_key: str) -> list[sqlite3.Row]:
    rows = con.execute(
        """
        SELECT * FROM cv_entries
        WHERE section_key = ?
          AND include_long = 1
        ORDER BY id
        """,
        (section_key,),
    ).fetchall()
    if section_key == "editorial_activities":
        return sorted(rows, key=lambda row: (0 if clean_cell(row["subcategory"]) == "Ad hoc Reviewer" else 1, row["id"]))
    return rows


def detail_columns(row: sqlite3.Row) -> tuple[str, str, str]:
    description = row_value(row, "description")
    parts = [clean_cell(part) for part in description.split("|")]
    parts = [part for part in parts if part]
    title = row_value(row, "title") or (parts[0] if parts else "")
    middle = ""
    organization = row_value(row, "organization") or row_value(row, "location") or ""
    if len(parts) >= 3:
        title = title or parts[0]
        middle = parts[1]
        organization = organization or parts[2]
    elif len(parts) == 2:
        middle = parts[1]
    elif row_value(row, "role"):
        middle = row_value(row, "role")
    details = row_value(row, "amount") or middle
    return title, details, organization


def trainee_achievement_map(con: sqlite3.Connection) -> dict[int, list[str]]:
    rows = con.execute(
        """
        SELECT t.cv_entry_id, a.title, a.organization, a.amount
        FROM trainee_achievements a
        JOIN trainees t ON t.id = a.trainee_id
        WHERE t.cv_entry_id IS NOT NULL
        ORDER BY a.year, a.id
        """
    ).fetchall()
    achievements: dict[int, list[str]] = {}
    for row in rows:
        title = row_value(row, "title")
        parts = [title] if title else []
        organization = row_value(row, "organization")
        amount = row_value(row, "amount")
        if organization:
            parts.append(organization)
        if amount:
            parts.append(amount)
        achievements.setdefault(row["cv_entry_id"], []).append(", ".join(parts))
    return achievements


def entries_table(
    rows: list[sqlite3.Row],
    achievements_by_entry: dict[int, list[str]] | None = None,
    include_amount: bool = False,
) -> list[str]:
    if include_amount:
        out = [
            f"| {tr('dates')} | {tr('role_title')} | {tr('funding_source')} | {tr('role')} | {tr('amount')} | {tr('details')} |",
            "| --- | --- | --- | --- | --- | --- |",
        ]
    else:
        out = [
            f"| {tr('dates')} | {tr('role_title')} | {tr('field_details')} | {tr('institution')} |",
            "| --- | --- | --- | --- |",
        ]
    for row in rows:
        title, details, organization = detail_columns(row)
        achievements = (achievements_by_entry or {}).get(row["id"], [])
        if achievements:
            suffix = f"{tr('achievements')}: " + "; ".join(achievements)
            details = f"{details}<br>{suffix}" if details else suffix
        if include_amount:
            out.append(
                f"| {md_escape(period(row['start_date'], row['end_date']))} | {md_escape(title)} | {md_escape(organization)} | {md_escape(row['role'])} | {md_escape(row['amount'])} | {md_escape(row['description'])} |"
            )
        else:
            out.append(
                f"| {md_escape(period(row['start_date'], row['end_date']))} | {md_escape(title)} | {md_escape(details)} | {md_escape(organization)} |"
            )
    return out


def publication_citation(row: sqlite3.Row, index: int) -> str:
    authors = row["authors"] or ""
    title = row["title"] or ""
    venue = row["venue"] or ""
    year = row["year"] or ""
    doi = row["doi"] or ""
    pmid = row["pmid"] or ""
    parts = []
    if authors:
        parts.append(authors)
    if title:
        parts.append(title)
    if venue:
        parts.append(f"*{venue}*")
    if year:
        parts.append(year)
    citation = ". ".join(parts).strip()
    extras = []
    if doi:
        extras.append(f"doi:{doi}")
    if pmid:
        extras.append(f"PMID:{pmid}")
    if extras:
        citation = f"{citation}. {'; '.join(extras)}"
    return f"{index}. {citation}"


def doi_url(row: sqlite3.Row) -> str:
    doi = clean_cell(row["doi"])
    if doi:
        return doi if doi.startswith("http") else f"https://doi.org/{doi}"
    return clean_cell(row["url"])


def markdown_publication_citation(row: sqlite3.Row) -> str:
    authors = clean_cell(row["authors"])
    title = clean_cell(row["title"])
    venue = clean_cell(row["venue"])
    year = clean_cell(row["year"])
    link = doi_url(row)
    authors = citation_cell(authors)
    title = citation_cell(title)
    venue = citation_cell(venue)
    year = citation_cell(year)
    authors = re.sub(r"\b(Andreas\s+Horn|Horn\s+A\.?)\b", r"**\1**", authors)
    parts = []
    if authors:
        parts.append(authors)
    if title:
        parts.append(title)
    if venue:
        parts.append(f"<u><em>{html_escape(venue.title() if venue.isupper() else venue)}</em></u>")
    if year:
        parts.append(year)
    impact = impact_factor_label(row)
    if impact:
        parts.append(impact)
    citation = ". ".join(parts).strip()
    if link:
        citation = f"{citation}. [{link}]({link})"
    return citation


def publications_block(con: sqlite3.Connection) -> list[str]:
    rows = con.execute(
        """
        SELECT * FROM publications
        WHERE COALESCE(suppress_display, 0) = 0
        ORDER BY
          CASE WHEN year IS NULL OR year = '' THEN 1 ELSE 0 END,
          CAST(year AS INTEGER),
          lower(title)
        """
    ).fetchall()
    out = [f"## {tr('scholarship')}", "", f"### {tr('peer_reviewed')}", ""]
    peer_rows = [row for row in rows if row["category"] == "peer_reviewed"]
    for index, row in enumerate(peer_rows, 1):
        out.append(f"{index}. {markdown_publication_citation(row)}")
    other_rows = [row for row in rows if row["category"] != "peer_reviewed"]
    if other_rows:
        out.extend(["", f"### {tr('other_scholarship')}", ""])
        for index, row in enumerate(other_rows, 1):
            out.append(f"{index}. {markdown_publication_citation(row)}")
    return out


def narrative_report_block(con: sqlite3.Connection) -> list[str]:
    row = con.execute("SELECT title, body, title_de, body_de FROM narrative_reports WHERE id=1").fetchone()
    if not row:
        return []
    title = row_value(row, "title") if not (LANG == "de" and not clean_cell(row["title_de"])) else tr("narrative_report")
    body = row_value(row, "body")
    if not clean_cell(body):
        return []
    return ["", f"## {title}", "", clean_cell(body)]


def build_markdown() -> str:
    con = connect()
    lines = person_block(con)
    lines.append("")
    achievements_by_entry = trainee_achievement_map(con)
    for section_key, title in SECTION_ORDER:
        rows = entry_rows(con, section_key)
        if not rows:
            continue
        lines.extend([f"## {localized_section_title(section_key, title)}", ""])
        lines.extend(
            entries_table(
                rows,
                achievements_by_entry if section_key == "mentoring" else None,
                include_amount=section_key == "funding",
            )
        )
        lines.append("")
    lines.extend(publications_block(con))
    lines.extend(narrative_report_block(con))
    con.close()
    return "\n".join(lines).rstrip() + "\n"


def typ_period(start: str | None, end: str | None) -> str:
    value = period(start, end)
    return value.replace("-", "-\n", 1) if "-" in value and len(value) > 9 else value


def row_period(row: sqlite3.Row) -> str:
    value = ""
    if row["raw_text"]:
        if "|" in row["raw_text"]:
            raw_parts = [part.strip() for part in row["raw_text"].split("|")]
            first = next((part for part in raw_parts if part), "")
        else:
            match = re.match(r"^(\d{1,2}/\d{1,2}/\d{2,4}-\d{1,2}/\d{1,2}/\d{2,4},?|\d{1,2}/\d{1,2}/\d{2,4}-?|\d{4}-\d{4}|\d{4}-?)\b", row["raw_text"].strip())
            first = match.group(1) if match else ""
        if looks_like_period(first):
            value = first
    if not value:
        value = period(row["start_date"], row["end_date"])
    return value.replace("-", "-\n", 1) if "-" in value and len(value) > 9 else value


def looks_like_period(value: str | None) -> bool:
    text = clean_cell(value)
    if not text:
        return False
    return bool(text[:4].isdigit() or "/" in text)


def typ_section(title: str) -> str:
    return f"\n#v(0.15in)\n{typ_text(title + ':', bold=True)}\n#v(0.055in)\n"


def typ_subheading(title: str) -> str:
    return f"#v(0.075in)\n{typ_text(title, bold=True)}\n#v(0.04in)\n"


def typ_grid(rows: list[list[str]], *, columns: str = "(1.05in, 1.42in, 2.05in, 1.7in)", row_gutter: str = "0.052in") -> str:
    cells = []
    for row in rows:
        for cell in row:
            cells.append(f"[{typ_text(cell)}]")
    return (
        f"#grid(columns: {columns}, gutter: 0.13in, row-gutter: {row_gutter},\n"
        + ",\n".join(f"  {cell}" for cell in cells)
        + "\n)\n"
    )


def typ_rich_text(value: str | None, *, bold_names: bool = False, underline: bool = False, italic: bool = False) -> str:
    text = clean_cell(value)
    if not text:
        return typ_text("")
    if bold_names:
        pieces = []
        pos = 0
        for match in re.finditer(r"\b(?:Andreas\s+Horn|Horn\s+A\.?)\b", text):
            if match.start() > pos:
                before = text[pos : match.start()]
                stripped = before.rstrip()
                if stripped:
                    pieces.append(typ_text(stripped))
                if before and before[-1].isspace():
                    pieces.append("#h(0.28em)")
            pieces.append(typ_text(match.group(0), bold=True))
            pos = match.end()
        if pos < len(text):
            after = text[pos:]
            if after and after[0].isspace():
                pieces.append("#h(0.28em)")
                after = after.lstrip()
            if after:
                pieces.append(typ_text(after))
        return "".join(pieces)
    body = typ_text(text)
    if italic:
        body = f"#emph[{body}]"
    if underline:
        body = f"#underline[{body}]"
    return body


def citation_cell(value: str | None) -> str:
    text = clean_cell(value).strip()
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r",(?=\S)", ", ", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.rstrip(" .")


def sentence_part(value: str | None) -> str:
    text = citation_cell(value)
    return text if not text or text.endswith((".", "?", "!")) else f"{text}."


def impact_factor_label(row: sqlite3.Row) -> str:
    value = row["impact_factor"] if "impact_factor" in row.keys() else None
    if value in (None, ""):
        return ""
    try:
        formatted = f"{float(value):g}"
    except (TypeError, ValueError):
        formatted = clean_cell(str(value))
    year = citation_cell(row["impact_factor_year"] if "impact_factor_year" in row.keys() else "")
    return f"IF {formatted} ({year})" if year else f"IF {formatted}"


def typ_link_text(url: str, display: str | None = None) -> str:
    return f'#link({typ_string(url)})[#underline[#text(fill: blue, {typ_string(display or url)})]]'


def normalize_url(url: str) -> str:
    if url.startswith(("http://", "https://")):
        return url
    return f"https://{url}"


def typ_text_with_links(value: str | None) -> str:
    text = clean_cell(value)
    if not text:
        return typ_text("")
    pieces = []
    pos = 0
    pattern = re.compile(r"(?<!@)\b(?:https?://[^\s)]+|www\.[^\s)]+)")
    for match in pattern.finditer(text):
        if match.start() > pos:
            pieces.append(typ_text(text[pos : match.start()]))
        display = match.group(0).rstrip(".,;")
        trailing = match.group(0)[len(display) :]
        pieces.append(typ_link_text(normalize_url(display), display))
        if trailing:
            pieces.append(typ_text(trailing))
        pos = match.end()
    if pos < len(text):
        pieces.append(typ_text(text[pos:]))
    return "".join(pieces)


def typ_publication_citation(row: sqlite3.Row) -> str:
    parts = []
    if row["authors"]:
        parts.append(typ_rich_text(sentence_part(row["authors"]), bold_names=True))
    if row["title"]:
        parts.append(typ_text(sentence_part(row["title"])))
    if row["venue"]:
        venue = citation_cell(row["venue"])
        venue = venue.title() if venue.isupper() else venue
        parts.append(typ_rich_text(sentence_part(venue), italic=True, underline=True))
    if row["year"]:
        parts.append(typ_text(sentence_part(row["year"])))
    impact = impact_factor_label(row)
    if impact:
        parts.append(typ_text(sentence_part(impact)))
    doi = citation_cell(row["doi"])
    if doi:
        parts.append(f'{typ_text("doi:")}{typ_link_text(doi_url(row), doi)}')
    else:
        link = doi_url(row)
        if link:
            parts.append(typ_link_text(link))
    return "#h(0.28em)".join(parts)


def typ_invited_presentations(rows: list[sqlite3.Row]) -> str:
    cells = []
    for row in rows:
        extra_parts = []
        title = row_value(row, "title")
        for text in [row_value(row, "organization"), row_value(row, "location"), row_value(row, "description")]:
            if text and text != title and text not in extra_parts:
                extra_parts.append(text)
        details = "\n".join(
            part
            for part in [
                title,
                "\n".join(extra_parts),
            ]
            if part
        )
        cells.append(f"[{typ_text(row_period(row))}]")
        cells.append(f"[{typ_text(details)}]")
    return (
        f"#block(below: 0.12in)[{typ_text(tr('no_sponsor'))}]\n"
        "#grid(columns: (0.56in, 6.04in), gutter: 0.22in, row-gutter: 0.07in,\n"
        + ",\n".join(f"  {cell}" for cell in cells)
        + "\n)\n"
    )


def typ_entry_rows(
    rows: list[sqlite3.Row],
    achievements_by_entry: dict[int, list[str]] | None = None,
    *,
    columns: int = 4,
) -> list[list[str]]:
    out = []
    for row in rows:
        title, details, organization = detail_columns(row)
        is_continuation = not clean_cell(row["start_date"]) and not clean_cell(row["end_date"])
        if row["section_key"] == "grant_review" and is_continuation and organization:
            period_title = row_period(row) or clean_cell(title)
            if out and period_title.replace("\n", "") in out[-1][0].replace("\n", ""):
                out[-1][-1] = "\n".join(part for part in [out[-1][-1], organization] if part)
                continue
        if is_continuation and (looks_like_period(title) or looks_like_period(row_period(row))) and organization:
            if out:
                out[-1][1] = "\n".join(part for part in [out[-1][1], row_period(row) or title] if part)
                out[-1][-1] = "\n".join(part for part in [out[-1][-1], organization] if part)
                continue
        achievements = (achievements_by_entry or {}).get(row["id"], [])
        if achievements:
            suffix = "Achievements: " + "; ".join(achievements)
            details = f"{details}\n{suffix}" if details else suffix
        if columns == 3:
            out.append([row_period(row), title, organization or details])
        else:
            out.append([row_period(row), title, details, organization])
    return out


def typ_funding_rows(rows: list[sqlite3.Row]) -> list[list[str]]:
    out = []
    for row in rows:
        out.append(
            [
                row_period(row),
                row_value(row, "title"),
                row_value(row, "organization"),
                "\n".join(part for part in [row_value(row, "role"), row_value(row, "amount")] if part),
            ]
        )
    return out


def typ_funding_blocks(rows: list[sqlite3.Row]) -> str:
    cells = []
    for row in rows:
        role_amount = row_value(row, "role")
        amount = row_value(row, "amount")
        if role_amount and amount:
            role_amount = f"{role_amount} ({amount})"
        elif amount:
            role_amount = amount
        details = "\n".join(
            part
            for part in [
                row_value(row, "title"),
                row_value(row, "organization"),
                role_amount,
                row_value(row, "description"),
            ]
            if part
        )
        cells.append(f"[{typ_text(row_period(row))}]")
        cells.append(f"[{typ_text(details)}]")
    return (
        "#grid(columns: (1.0in, 5.62in), gutter: 0.13in, row-gutter: 0.095in,\n"
        + ",\n".join(f"  {cell}" for cell in cells)
        + "\n)\n"
    )


def typ_editorial_blocks(rows: list[sqlite3.Row]) -> list[str]:
    lines = []
    reviewer_rows = [row for row in rows if clean_cell(row["subcategory"]) == "Ad hoc Reviewer"]
    role_rows = [row for row in rows if clean_cell(row["subcategory"]) != "Ad hoc Reviewer"]
    for row in reviewer_rows:
        lines.append(typ_subheading(tr("ad_hoc_reviewer")))
        reviewer_text = row_value(row, "description") or row_value(row, "raw_text")
        lines.append(f"#block(below: 0.08in)[{typ_text(reviewer_text)}]\n")
    if role_rows:
        lines.append(typ_subheading(tr("other_editorial_roles")))
        lines.append(typ_grid(typ_entry_rows(role_rows, columns=3), columns="(1.0in, 3.1in, 2.5in)", row_gutter="0.075in"))
    return lines


def typ_mentoring_blocks(rows: list[sqlite3.Row], achievements_by_entry: dict[int, list[str]]) -> list[str]:
    lines = []
    general_rows = [row for row in rows if clean_cell(row["title"]).startswith("Supervision of PhD students")]
    trainee_rows = [row for row in rows if row not in general_rows]
    if general_rows:
        lines.append(typ_grid(typ_entry_rows(general_rows, columns=3), columns="(1.0in, 3.1in, 2.5in)", row_gutter="0.075in"))
    if trainee_rows:
        lines.append(typ_subheading(tr("other_trainees")))
        cells = []
        for row in trainee_rows:
            details = row_value(row, "title")
            achievements = achievements_by_entry.get(row["id"], [])
            if achievements:
                details = f"{details}\n{tr('achievements')}: " + "; ".join(achievements)
            cells.append(f"[{typ_text(row_period(row))}]")
            cells.append(f"[{typ_text(details)}]")
        lines.append(
            "#grid(columns: (1.0in, 5.62in), gutter: 0.13in, row-gutter: 0.085in,\n"
            + ",\n".join(f"  {cell}" for cell in cells)
            + "\n)\n"
        )
    return lines


def grouped_rows(rows: list[sqlite3.Row]) -> list[tuple[str | None, list[sqlite3.Row]]]:
    groups: list[tuple[str | None, list[sqlite3.Row]]] = []
    for row in rows:
        subcategory = row_value(row, "subcategory") or clean_cell(row["subcategory"]) or None
        if not groups or groups[-1][0] != subcategory:
            groups.append((subcategory, []))
        groups[-1][1].append(row)
    return groups


def build_typst() -> str:
    con = connect()
    person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
    today = dt.datetime.now().strftime("%d.%m.%Y") if LANG == "de" else dt.datetime.now().strftime("%B %-d, %Y")
    achievements_by_entry = trainee_achievement_map(con)

    lines = [
        '#set page(width: 8.5in, height: 11in, margin: (left: 0.73in, right: 0.62in, top: 0.72in, bottom: 0.55in))',
        f'#set text(font: "Helvetica", size: 10.5pt, lang: "{LANG}")',
        "#set par(leading: 0.49em)",
        "#align(center)[",
        f"  {typ_text(tr('faculty'), bold=True)}",
        "  #linebreak()",
        f"  {typ_text(tr('cv'), bold=True)}",
        "]",
        "#v(0.32in)",
    ]
    if person:
        metadata = [
            (tr("date_prepared"), today),
            (tr("name"), person["display_name"] or person["full_name"]),
            (tr("office_address"), person["office_address"]),
            (tr("home_address"), person["home_address"]),
            (tr("work_phone"), person["work_phone"]),
            (tr("work_email"), person["work_email"]),
            (tr("place_of_birth"), person["place_of_birth"]),
        ]
        cells = []
        for label, value in metadata:
            if value:
                cells.append(f"[{typ_text(label, bold=True)}]")
                cells.append(f"[{typ_text(value, bold=True)}]")
        lines.append("#grid(columns: (1.38in, 4.9in), row-gutter: 0.11in,\n" + ",\n".join(f"  {cell}" for cell in cells) + "\n)")

    for section_key, title in SECTION_ORDER:
        rows = entry_rows(con, section_key)
        if not rows:
            continue
        lines.append(typ_section(localized_section_title(section_key, title)))
        if section_key == "funding":
            lines.append(typ_funding_blocks(rows))
        elif section_key == "editorial_activities":
            lines.extend(typ_editorial_blocks(rows))
        elif section_key == "mentoring":
            lines.extend(typ_mentoring_blocks(rows, achievements_by_entry))
        elif section_key == "invited_presentations":
            lines.append(typ_invited_presentations(rows))
        elif section_key in THREE_COLUMN_SECTIONS:
            for subcategory, grouped in grouped_rows(rows):
                if subcategory:
                    lines.append(typ_subheading(subcategory))
                lines.append(typ_grid(typ_entry_rows(grouped, columns=3), columns="(1.0in, 2.55in, 3.02in)", row_gutter="0.075in"))
        else:
            lines.append(typ_grid(typ_entry_rows(rows, achievements_by_entry if section_key == "mentoring" else None)))

    pub_rows = con.execute(
        """
        SELECT * FROM publications
        WHERE COALESCE(suppress_display, 0) = 0
        ORDER BY
          CASE WHEN year IS NULL OR year = '' THEN 1 ELSE 0 END,
          CAST(year AS INTEGER),
          lower(title)
        """
    ).fetchall()
    if pub_rows:
        lines.append(typ_section(tr("scholarship")))
        lines.append(typ_text(tr("peer_reviewed"), bold=True))
        lines.append("#v(0.045in)")
        for index, row in enumerate([row for row in pub_rows if row["category"] == "peer_reviewed"], 1):
            lines.append(
                "#grid(columns: (0.34in, 6.28in), gutter: 0.08in, row-gutter: 0pt,\n"
                f"  [{typ_text(str(index) + '.')}],\n"
                f"  [{typ_publication_citation(row)}]\n"
                ")"
            )
        other_rows = [row for row in pub_rows if row["category"] != "peer_reviewed"]
        if other_rows:
            lines.append(typ_subheading(tr("other_scholarship")))
            for index, row in enumerate(other_rows, 1):
                lines.append(
                    "#grid(columns: (0.34in, 6.28in), gutter: 0.08in, row-gutter: 0pt,\n"
                    f"  [{typ_text(str(index) + '.')}],\n"
                    f"  [{typ_publication_citation(row)}]\n"
                    ")"
                )
    report = con.execute("SELECT title, body, title_de, body_de FROM narrative_reports WHERE id=1").fetchone()
    if report and row_value(report, "body"):
        report_title = row_value(report, "title") if not (LANG == "de" and not clean_cell(report["title_de"])) else tr("narrative_report")
        lines.append(typ_section(report_title or tr("narrative_report")))
        lines.append(f"#block[#set par(justify: true)\n{typ_text_with_links(row_value(report, 'body'))}]")
    con.close()
    return "\n".join(lines) + "\n"


def markdown_to_html(markdown: str) -> tuple[str, str | None]:
    body, warning = markdown_to_html_body(markdown, ROOT)
    warning_html = f'<p class="warning">{html.escape(warning)}</p>' if warning else ""
    html_doc = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Long CV Preview</title>
  <style>
    @page {{ size: letter; margin: 0.65in; }}
    body {{ font-family: Arial, Helvetica, sans-serif; margin: 28px auto; max-width: 980px; line-height: 1.28; color: #111; font-size: 11px; }}
    .cv-kicker {{ font-weight: 700; text-align: center; margin-bottom: 8px; }}
    h1 {{ font-size: 18px; text-align: center; margin: 6px 0 14px; }}
    h2 {{ font-size: 13px; margin: 20px 0 6px; font-weight: 700; }}
    h3 {{ font-size: 11px; margin: 14px 0 6px; font-weight: 700; }}
    table {{ border-collapse: collapse; width: 100%; margin: 4px 0 12px; font-size: 10px; page-break-inside: auto; }}
    tr {{ page-break-inside: avoid; page-break-after: auto; }}
    th, td {{ border: 1px solid #9b9b9b; padding: 4px 5px; vertical-align: top; }}
    th {{ background: #f1f1f1; text-align: left; font-weight: 700; }}
    td:first-child, th:first-child {{ width: 112px; }}
    p {{ margin: 0 0 6px; }}
    .warning {{ border: 1px solid #d8b24c; background: #fff8db; padding: 8px; color: #5d4700; }}
    @media print {{
      body {{ margin: 0; max-width: none; }}
      a {{ color: #111; text-decoration: none; }}
    }}
  </style>
</head>
<body>
{warning_html}
{body}
</body>
</html>
"""
    return html_doc, warning


def output_stem() -> str:
    return "long_cv_de" if LANG == "de" else "long_cv"


def set_run_font(run, *, size: float = 11, bold: bool | None = None, italic: bool | None = None, underline: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 0, line_spacing: float = 1.0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_paragraph_border(paragraph, *, top: bool = False, bottom: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge, enabled in (("top", top), ("bottom", bottom)):
        if not enabled:
            continue
        node = p_bdr.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            p_bdr.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), "000000")


def add_text_paragraph(doc: Document, text_value: str = "", *, bold: bool = False, size: float = 11, before: float = 0, after: float = 0, justify: bool = False):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text_value:
        run = paragraph.add_run(text_value)
        set_run_font(run, size=size, bold=bold)
    return paragraph


def add_section_heading(doc: Document, marker: str, title: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=8, after=2)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.35))
    run = paragraph.add_run(f"{marker}\t{title}")
    set_run_font(run, size=11, bold=True)


def add_subheading(doc: Document, title: str) -> None:
    paragraph = add_text_paragraph(doc, title, bold=True, before=3, after=1)
    for run in paragraph.runs:
        set_run_font(run, size=11, bold=True)


def add_tabbed_paragraph(doc: Document, left: str, right: str, *, left_width: float = 1.28, size: float = 11) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(left_width))
    run = paragraph.add_run(clean_cell(left))
    set_run_font(run, size=size)
    run = paragraph.add_run("\t" + clean_cell(right))
    set_run_font(run, size=size)


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for label, value in rows:
        if not clean_cell(value):
            continue
        cells = table.add_row().cells
        cells[0].text = f"{label}:"
        cells[1].text = clean_cell(value)
        cells[0].width = Inches(1.55)
        cells[1].width = Inches(5.75)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph)
                for run in paragraph.runs:
                    set_run_font(run, size=11, bold=(cell is cells[0]))
    table.style = "Table Grid"


def format_entry_text(row: sqlite3.Row, *, include_amount: bool = False) -> str:
    title, details, organization = detail_columns(row)
    parts = [title]
    if include_amount:
        parts = [row_value(row, "title"), row_value(row, "organization"), row_value(row, "role"), row_value(row, "amount"), row_value(row, "description")]
    else:
        parts.extend([details, organization])
    return "; ".join(clean_cell(part) for part in parts if clean_cell(part))


def add_entry_rows_docx(doc: Document, rows: list[sqlite3.Row], achievements_by_entry: dict[int, list[str]] | None = None, *, include_amount: bool = False) -> None:
    for row in rows:
        details = format_entry_text(row, include_amount=include_amount)
        achievements = (achievements_by_entry or {}).get(row["id"], [])
        if achievements:
            details = f"{details}; {tr('achievements')}: " + "; ".join(achievements)
        add_tabbed_paragraph(doc, row_period(row).replace("\n", "-"), details)


def add_publication_docx(doc: Document, index: int, row: sqlite3.Row) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.32))
    run = paragraph.add_run(f"{index}.\t")
    set_run_font(run, size=10.5)
    authors = citation_cell(row["authors"])
    title = citation_cell(row["title"])
    venue = citation_cell(row["venue"])
    year = citation_cell(row["year"])
    doi = citation_cell(row["doi"])
    impact = impact_factor_label(row)
    pieces = [(sentence_part(authors), False, False), (sentence_part(title), False, False), (sentence_part(venue.title() if venue.isupper() else venue), True, True), (sentence_part(year), False, False), (sentence_part(impact), False, False)]
    first = True
    for text_value, italic, underline in pieces:
        if not text_value:
            continue
        if not first:
            sep = paragraph.add_run(" ")
            set_run_font(sep, size=10.5)
        first = False
        pos = 0
        for match in re.finditer(r"\b(?:Andreas\s+Horn|Horn\s+A\.?)\b", text_value):
            if match.start() > pos:
                run = paragraph.add_run(text_value[pos:match.start()])
                set_run_font(run, size=10.5, italic=italic, underline=underline)
            run = paragraph.add_run(match.group(0))
            set_run_font(run, size=10.5, bold=True, italic=italic, underline=underline)
            pos = match.end()
        if pos < len(text_value):
            run = paragraph.add_run(text_value[pos:])
            set_run_font(run, size=10.5, italic=italic, underline=underline)
    if doi:
        run = paragraph.add_run(f" doi:{doi}")
        set_run_font(run, size=10.5)


def add_wrapped_body_paragraphs(doc: Document, value: str) -> None:
    for block in re.split(r"\n\s*\n", clean_cell(value)):
        text_value = " ".join(line.strip() for line in block.splitlines() if line.strip())
        if text_value:
            add_text_paragraph(doc, text_value, size=11, after=2, justify=True)


def build_docx(path: Path, lang: str = "en") -> Path:
    global LANG
    LANG = "de" if lang == "de" else "en"
    con = connect()
    person = con.execute("SELECT * FROM person WHERE id=1").fetchone()
    today = dt.datetime.now().strftime("%d.%m.%Y") if LANG == "de" else dt.datetime.now().strftime("%B %-d, %Y")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=6)
    add_paragraph_border(title, top=True, bottom=True)
    run = title.add_run(tr("cv").upper())
    set_run_font(run, size=13, bold=True)

    if person:
        rows = [
            (tr("date_prepared").rstrip(":"), today),
            (tr("name").rstrip(":"), person["display_name"] or person["full_name"]),
            (tr("office_address").rstrip(":"), person["office_address"]),
            (tr("work_phone").rstrip(":"), person["work_phone"]),
            (tr("work_email").rstrip(":"), person["work_email"]),
            (tr("place_of_birth").rstrip(":"), person["place_of_birth"]),
        ]
        add_metadata_table(doc, rows)

    achievements_by_entry = trainee_achievement_map(con)
    marker_ord = ord("A")
    for section_key, section_title in SECTION_ORDER:
        rows = entry_rows(con, section_key)
        if not rows:
            continue
        add_section_heading(doc, f"{chr(marker_ord)}.", localized_section_title(section_key, section_title))
        marker_ord += 1
        if section_key in {"editorial_activities", "mentoring"}:
            for subcategory, grouped in grouped_rows(rows):
                if subcategory:
                    add_subheading(doc, subcategory)
                add_entry_rows_docx(
                    doc,
                    grouped,
                    achievements_by_entry if section_key == "mentoring" else None,
                    include_amount=False,
                )
        else:
            add_entry_rows_docx(
                doc,
                rows,
                achievements_by_entry if section_key == "mentoring" else None,
                include_amount=section_key == "funding",
            )

    publication_rows = con.execute(
        """
        SELECT * FROM publications
        WHERE COALESCE(suppress_display, 0) = 0
        ORDER BY
          CASE WHEN year IS NULL OR year = '' THEN 1 ELSE 0 END,
          CAST(year AS INTEGER),
          lower(title)
        """
    ).fetchall()
    peer_rows = [row for row in publication_rows if row["category"] == "peer_reviewed"]
    other_rows = [row for row in publication_rows if row["category"] != "peer_reviewed"]
    if peer_rows or other_rows:
        add_section_heading(doc, f"{chr(marker_ord)}.", tr("scholarship"))
        marker_ord += 1
        if peer_rows:
            add_subheading(doc, tr("peer_reviewed"))
            for index, row in enumerate(peer_rows, 1):
                add_publication_docx(doc, index, row)
        if other_rows:
            add_subheading(doc, tr("other_scholarship"))
            for index, row in enumerate(other_rows, 1):
                add_publication_docx(doc, index, row)

    report = con.execute("SELECT title, body, title_de, body_de FROM narrative_reports WHERE id=1").fetchone()
    if report and clean_cell(row_value(report, "body")):
        add_section_heading(doc, f"{chr(marker_ord)}.", row_value(report, "title") or tr("narrative_report"))
        add_wrapped_body_paragraphs(doc, row_value(report, "body"))

    con.close()
    if person:
        doc.core_properties.author = person["display_name"] or person["full_name"] or ""
    doc.core_properties.title = tr("cv")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build(lang: str = "en") -> dict[str, str]:
    global LANG
    LANG = "de" if lang == "de" else "en"
    OUTPUT.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown()
    stem = output_stem()
    md_path = OUTPUT / f"{stem}.md"
    html_path = OUTPUT / f"{stem}.html"
    docx_path = OUTPUT / f"{stem}.docx"
    pdf_path = OUTPUT / f"{stem}.pdf"
    typ_path = OUTPUT / f"{stem}.typ"
    md_path.write_text(markdown, encoding="utf-8")
    html_doc, html_warning = markdown_to_html(markdown)
    html_path.write_text(html_doc, encoding="utf-8")
    typ_path.write_text(build_typst(), encoding="utf-8")
    pdf, warning = compile_typst_if_available(typ_path, pdf_path, ROOT)
    docx = build_docx(docx_path, LANG)
    result = {
        "markdown": f"output/{output_ref(md_path)}",
        "html": f"output/{output_ref(html_path)}",
        "typst": f"output/{output_ref(typ_path)}",
    }
    if pdf:
        result["pdf"] = f"output/{output_ref(pdf)}"
    if docx:
        result["docx"] = f"output/{output_ref(docx)}"
    warnings = [item for item in (html_warning, warning) if item]
    if warnings:
        result["warning"] = " ".join(warnings)
    return result


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--lang", choices=["en", "de"], default="en")
    args = parser.parse_args()
    for name, path in build(args.lang).items():
        print(f"{name}: {path}")
